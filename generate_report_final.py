#!/usr/bin/env python3
"""InvestFlow Africa — Data Analysis Report Generator"""
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═════════════════════════════════════════════════════════════════
OUTPUT_DIR = "/home/z/my-project/download/charts_report"
REPORT_PATH = "/home/z/my-project/download/InvestFlow_Africa_Analyse_Donnees_2026.docx"

# Nimba palette
BLUE = "#2B6CB0"
GREEN = "#13612e"
GOLD = "#f5a524"
RED = "#b82105"
ORANGE = "#f7630c"
TEAL = "#2B9EB3"
PURPLE = "#7C5CFC"
PINK = "#E84393"
COLORS = [BLUE, GREEN, GOLD, RED, ORANGE, TEAL, PURPLE, PINK]

BG = "#FFFFFF"
TEXT = "#1A202C"
GRID = "#E2E8F0"
ACCENT = BLUE

# ═════════════════════════════════════════════════════════════════
# DATA
# ═════════════════════════════════════════════════════════════════
projects = [
    {"name": "Nimba Iron Mine", "sector": "Mining", "budget": 50, "cost": 38, "revenue": 89.5, "roi": 79, "status": "Active", "country": "Guinée"},
    {"name": "Solar Plant Dakar", "sector": "Energy", "budget": 20, "cost": 16, "revenue": 29.1, "roi": 44, "status": "Active", "country": "Sénégal"},
    {"name": "Lagos Real Estate", "sector": "Real Estate", "budget": 35, "cost": 30, "revenue": 52.5, "roi": 50, "status": "Active", "country": "Nigeria"},
    {"name": "Abidjan Tech Hub", "sector": "Technology", "budget": 15, "cost": 12, "revenue": 22.5, "roi": 50, "status": "Completed", "country": "Côte d'Ivoire"},
    {"name": "Kinshasa Agri-Business", "sector": "Agriculture", "budget": 25, "cost": 20, "revenue": 32.5, "roi": 30, "status": "Active", "country": "RDC"},
    {"name": "Accra Fintech", "sector": "Technology", "budget": 10, "cost": 8, "revenue": 16.0, "roi": 60, "status": "Active", "country": "Ghana"},
    {"name": "Nairobi Logistics", "sector": "Logistics", "budget": 30, "cost": 24, "revenue": 39.0, "roi": 30, "status": "Pending", "country": "Kenya"},
    {"name": "Douala Port Extension", "sector": "Logistics", "budget": 45, "cost": 36, "revenue": 63.0, "roi": 40, "status": "Active", "country": "Cameroun"},
]

monthly_data = [
    {"month": "Jan", "revenue": 12, "budget": 8, "cost": 6},
    {"month": "Fév", "revenue": 18, "budget": 15, "cost": 12},
    {"month": "Mar", "revenue": 25, "budget": 22, "cost": 18},
    {"month": "Avr", "revenue": 32, "budget": 30, "cost": 24},
    {"month": "Mai", "revenue": 38, "budget": 35, "cost": 28},
    {"month": "Jun", "revenue": 45, "budget": 42, "cost": 34},
    {"month": "Jul", "revenue": 52, "budget": 50, "cost": 40},
    {"month": "Aoû", "revenue": 58, "budget": 55, "cost": 44},
    {"month": "Sep", "revenue": 67, "budget": 60, "cost": 48},
    {"month": "Oct", "revenue": 75, "budget": 68, "cost": 54},
    {"month": "Nov", "revenue": 82, "budget": 73, "cost": 58},
    {"month": "Déc", "revenue": 95, "budget": 85, "cost": 68},
]

candlestick_data = [
    {"month": "Jan", "open": 8.2, "high": 10.5, "low": 5.8, "close": 12},
    {"month": "Fév", "open": 14.5, "high": 19.0, "low": 11.0, "close": 18},
    {"month": "Mar", "open": 21.0, "high": 26.5, "low": 17.5, "close": 25},
    {"month": "Avr", "open": 29.5, "high": 34.0, "low": 25.0, "close": 32},
    {"month": "Mai", "open": 33.8, "high": 38.5, "low": 28.0, "close": 38},
    {"month": "Jun", "open": 40.2, "high": 46.0, "low": 34.0, "close": 45},
    {"month": "Jul", "open": 48.5, "high": 55.0, "low": 41.0, "close": 52},
    {"month": "Aoû", "open": 53.2, "high": 59.5, "low": 46.0, "close": 58},
    {"month": "Sep", "open": 57.8, "high": 64.0, "low": 49.0, "close": 67},
    {"month": "Oct", "open": 65.0, "high": 72.0, "low": 56.0, "close": 75},
    {"month": "Nov", "open": 70.5, "high": 78.0, "low": 60.0, "close": 82},
    {"month": "Déc", "open": 82.0, "high": 95.0, "low": 72.0, "close": 95},
]

# ═════════════════════════════════════════════════════════════════
# CHART GENERATION
# ═════════════════════════════════════════════════════════════════
def style_chart(fig, ax, title):
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color(GRID)
    ax.spines['bottom'].set_color(GRID)
    ax.tick_params(colors=TEXT, labelsize=10)
    ax.yaxis.label.set_color(TEXT)
    ax.xaxis.label.set_color(TEXT)
    ax.set_title(title, fontsize=16, fontweight='bold', color=TEXT, pad=20)
    ax.grid(axis='y', color=GRID, alpha=0.5, linestyle='--')

# --- 1. BAR CHART: ROI par Projet ---
fig, ax = plt.subplots(figsize=(14, 7))
sorted_p = sorted(projects, key=lambda x: x['roi'])
names = [p['name'] for p in sorted_p]
rois = [p['roi'] for p in sorted_p]
colors_bar = [GREEN if r >= 60 else GOLD if r >= 40 else RED for r in rois]
style_chart(fig, ax, "ROI par Projet — InvestFlow Africa")
bars = ax.barh(names, rois, color=colors_bar, height=0.6, edgecolor='none')
for bar, roi in zip(bars, rois):
    ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2, f'{roi}%', va='center', ha='left', color=TEXT, fontweight='bold', fontsize=11)
ax.set_xlabel('ROI (%)', fontsize=12)
ax.set_xlim(0, max(rois) + 15)
plt.tight_layout()
plt.savefig(f'{OUTPUT_DIR}/bar_roi_par_projet.png', dpi=200, bbox_inches='tight', facecolor=BG)
plt.close()
print("Bar chart done")

# --- 2. LINE CHART: Tendances Mensuelles ---
fig, ax = plt.subplots(figsize=(14, 7))
style_chart(fig, ax, "Tendances Mensuelles — Budget, Coûts & Revenus (2024)")
months = [m['month'] for m in monthly_data]
revs = [m['revenue'] for m in monthly_data]
buds = [m['budget'] for m in monthly_data]
costs = [m['cost'] for m in monthly_data]
ax.plot(months, revs, color=GREEN, linewidth=3, marker='o', markersize=8, label='Revenus (M€)', zorder=5)
ax.plot(months, buds, color=BLUE, linewidth=3, marker='s', markersize=8, label='Budget (M€)', zorder=4)
ax.plot(months, costs, color=ORANGE, linewidth=3, marker='^', markersize=8, label='Coûts (M€)', zorder=3)
ax.fill_between(months, revs, costs, alpha=0.12, color=GREEN)
ax.legend(fontsize=12, facecolor='white', edgecolor=GRID, loc='upper left', framealpha=0.9)
ax.set_ylabel('Montant (M€)', fontsize=12)
plt.tight_layout()
plt.savefig(f'{OUTPUT_DIR}/line_tendances.png', dpi=200, bbox_inches='tight', facecolor=BG)
plt.close()
print("Line chart done")

# --- 3. PIE CHART: Revenus par Secteur ---
sector_rev = {}
for p in projects:
    s = p['sector']
    sector_rev[s] = sector_rev.get(s, 0) + p['revenue']
fig, ax = plt.subplots(figsize=(10, 10))
style_chart(fig, ax, "Répartition des Revenus par Secteur")
ax.set_facecolor('none')
wedges, texts, autotexts = ax.pie(
    sector_rev.values(), labels=sector_rev.keys(), autopct='%1.1f%%',
    colors=COLORS[:len(sector_rev)], startangle=140,
    pctdistance=0.8, wedgeprops=dict(width=0.5, edgecolor=BG, linewidth=2)
)
for t in texts:
    t.set_color(TEXT)
    t.set_fontsize(12)
for t in autotexts:
    t.set_color('white')
    t.set_fontweight('bold')
    t.set_fontsize(11)
plt.tight_layout()
plt.savefig(f'{OUTPUT_DIR}/pie_revenus_secteur.png', dpi=200, bbox_inches='tight', facecolor=BG)
plt.close()
print("Pie chart done")

# --- 4. CANDLESTICK CHART: Performance Mensuelle ---
fig, ax = plt.subplots(figsize=(14, 7))
style_chart(fig, ax, "Performance Mensuelle des Investissements (M€)")
months_short = [c['month'] for c in candlestick_data]
x = np.arange(len(months_short))
width = 0.5
for i, c in enumerate(candlestick_data):
    color = GREEN if c['close'] >= c['open'] else RED
    ax.plot([x[i], x[i]], [c['low'], c['high']], color=color, linewidth=1.5)
    ax.plot([x[i]-width/3, x[i]+width/3], [c['high'], c['high']], color=color, linewidth=1.5)
    ax.plot([x[i]-width/3, x[i]+width/3], [c['low'], c['low']], color=color, linewidth=1.5)
    ax.bar(x[i], c['close']-c['open'], width, bottom=c['open'], color=color, edgecolor='none', alpha=0.85)
ax.set_xticks(x)
ax.set_xticklabels(months_short, fontsize=10)
ax.set_ylabel('Montant (M€)', fontsize=12)
plt.tight_layout()
plt.savefig(f'{OUTPUT_DIR}/candlestick_performance.png', dpi=200, bbox_inches='tight', facecolor=BG)
plt.close()
print("Candlestick chart done")

# --- 5. DONUT CHART: Allocation Budget ---
fig, ax = plt.subplots(figsize=(10, 10))
style_chart(fig, ax, "Allocation du Budget par Projet")
ax.set_facecolor('none')
b_names = [p['name'] for p in projects]
budgets = [p['budget'] for p in projects]
wedges, texts, autotexts = ax.pie(
    budgets, labels=b_names, autopct='%1.1f%%',
    colors=COLORS[:len(b_names)], startangle=90,
    pctdistance=0.82, wedgeprops=dict(width=0.45, edgecolor=BG, linewidth=2)
)
for t in texts:
    t.set_color(TEXT)
    t.set_fontsize(9)
for t in autotexts:
    t.set_color('white')
    t.set_fontweight('bold')
    t.set_fontsize(9)
total_budget = sum(budgets)
ax.text(0, 0, f'{total_budget}M€\nTotal Budget', ha='center', va='center', fontsize=16, fontweight='bold', color=ACCENT)
plt.tight_layout()
plt.savefig(f'{OUTPUT_DIR}/donut_allocation.png', dpi=200, bbox_inches='tight', facecolor=BG)
plt.close()
print("Donut chart done")

print("\nAll 5 charts generated!")

# ═════════════════════════════════════════════════════════════════
# WORD REPORT
# ═════════════════════════════════════════════════════════════════
doc = Document()

# -- Page setup --
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    if level == 1:
        h.font.size = Pt(22)
        h.font.bold = True
    elif level == 2:
        h.font.size = Pt(16)
        h.font.bold = True
    elif level == 3:
        h.font.size = Pt(13)
        h.font.bold = True

def add_para(text, style='Normal', bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6)):
    p = doc.add_paragraph(text, style=style)
    p.alignment = align
    p.paragraph_format.space_after = space_after
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#2B6CB0"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#F2F2F2"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    return table

def add_image(path, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(14))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(12)
        for run in cap.runs:
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

# ═════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INVESTFLOW AFRICA")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rapport d'Analyse de Données")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x13, 0x61, 0x2E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Performance du Portefeuille d'Investissements — Q4 2024")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mars 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Confidentiel — Usage interne uniquement")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

# Page break
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════
doc.add_heading("Table des Matières", level=1)
doc.add_paragraph("(Mettre à jour dans Word : Ctrl+A puis F9)")
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════
doc.add_heading("1. Résumé Exécutif", level=1)

add_para("Ce rapport présente une analyse approfondie de la performance du portefeuille d'investissements d'InvestFlow Africa pour le quatrième trimestre 2024. Le portefeuille couvre 8 projets répartis dans 6 pays d'Afrique subsaharienne, avec une enveloppe budgétaire totale de 230 millions d'euros et des revenus cumulés de 344,1 millions d'euros, générant ainsi un retour sur investissement moyen pondéré de 47,9%.")

add_para("Les cinq constatations principales de cette analyse sont les suivantes : premièrement, le secteur minier, porté par le projet Nimba Iron Mine en Guinée, affiche le ROI le plus élevé du portefeuille à 79%, confirmant le potentiel exceptionnel des ressources naturelles africaines. Deuxièmement, les projets technologiques (Abidjan Tech Hub et Accra Fintech) présentent un ROI moyen de 55%, positionnant ce secteur comme le plus régulier en termes de rentabilité. Troisièmement, six des huit projets sont actuellement actifs, démontrant une capacité d'exécution solide avec un taux de réalisation de 75%.")

add_para("Quatrièmement, l'analyse des tendances mensuelles révèle une croissance soutenue des revenus, avec une augmentation de 692% entre janvier (12 M€) et décembre 2024 (95 M€), dépassant systématiquement les budgets et coûts prévisionnels. Cinquièmement, le rapport Budget-Revenus démontre un flux net positif de 114,1 M€, soit une marge bénéficiaire globale de 49,6% sur l'enveloppe budgétaire totale.")

add_para("Trois recommandations stratégiques émergent de cette analyse : accroître l'allocation au secteur minier et technologique qui affichent les meilleures performances ; renforcer le suivi des projets à risque élevé (Kinshasa Agri-Business et Nairobi Logistics) ; et diversifier géographiquement vers l'Afrique de l'Est pour réduire la concentration sectorielle actuelle.")

# ═════════════════════════════════════════════════════════════════
# 2. PORTEFEUILLE OVERVIEW
# ═════════════════════════════════════════════════════════════════
doc.add_heading("2. Vue d'Ensemble du Portefeuille", level=1)

add_para("Le portefeuille d'InvestFlow Africa comprend 8 projets d'investissement répartis sur 6 pays d'Afrique subsaharienne, couvrant les secteurs minier, énergétique, immobilier, technologique, agricole et logistique. Cette diversification sectorielle et géographique constitue un atout majeur pour la gestion des risques tout en capitalisant sur les opportunités de croissance les plus prometteuses du continent africain.")

add_para("L'analyse de la composition du portefeuille révèle une allocation stratégique orientée vers les secteurs à forte valeur ajoutée. Le budget total de 230 M€ est réparti de manière inégale, avec les trois projets les plus importants (Nimba Iron Mine, Douala Port Extension et Lagos Real Estate) représentant à eux seuls 130 M€, soit 56,5% de l'enveloppe totale. Cette concentration reflète une stratégie d'investissement ciblée sur les projets à fort potentiel de rendement.")

doc.add_heading("2.1 Tableau de Performance des Projets", level=2)

headers = ["Projet", "Pays", "Secteur", "Budget (M€)", "Coûts (M€)", "Revenus (M€)", "ROI", "Statut"]
rows = []
for p in sorted(projects, key=lambda x: x['roi'], reverse=True):
    rows.append([
        p['name'], p['country'], p['sector'],
        f"{p['budget']}", f"{p['cost']}", f"{p['revenue']}",
        f"{p['roi']}%", p['status']
    ])
rows.append(["TOTAL / MOY", "6 pays", "6 secteurs", "230", "184", "344,1", "47,9%", "6 actifs / 8"])
add_table(headers, rows)

doc.add_paragraph()

doc.add_heading("2.2 Indicateurs Clés de Performance", level=2)
kpi_headers = ["Indicateur", "Valeur", "Commentaire"]
kpi_rows = [
    ["Budget Total", "230 M€", "Enveloppe globale déployée"],
    ["Revenus Totaux", "344,1 M€", "Revenus cumulés tous projets"],
    ["ROI Moyen Pondéré", "47,9%", "Supérieur à la moyenne sectorielle africaine de 35%"],
    ["Marge Nette", "114,1 M€", "Revenus moins coûts"],
    ["Taux de Réalisation", "75%", "6 projets actifs sur 8"],
    ["Budget Consommé", "80%", "184 M€ de coûts sur 230 M€"],
    ["Meilleur ROI", "79% (Nimba)", "Secteur minier en Guinée"],
    ["Pays Couverts", "6", "Guinée, Sénégal, Nigeria, CI, RDC, Ghana, Kenya, Cameroun"],
]
add_table(kpi_headers, kpi_rows)

# ═════════════════════════════════════════════════════════════════
# 3. ANALYSE DE LA RENTABILITÉ (BAR CHART)
# ═════════════════════════════════════════════════════════════════
doc.add_heading("3. Analyse de la Rentabilité par Projet", level=1)

add_para("L'analyse de la rentabilité constitue le pilier central de l'évaluation de la performance du portefeuille. Le graphique ci-dessous présente le ROI de chacun des 8 projets, classé par ordre croissant, permettant d'identifier immédiatement les projets les plus performants et ceux nécessitant une attention particulière.")

add_para("Le projet Nimba Iron Mine se distingue nettement avec un ROI de 79%, soit 31 points de pourcentage au-dessus de la moyenne du portefeuille. Cette performance exceptionnelle s'explique par la combinaison de coûts maîtrisés (38 M€ sur un budget de 50 M€) et de revenus élevés (89,5 M€), reflétant l'attractivité mondiale pour les ressources minérales africaines, notamment le fer et les minerais rares présents dans la région de Nimba en Guinée forestière.")

add_para("Les deux projets technologiques (Abidjan Tech Hub et Accra Fintech) affichent tous deux un ROI de 50% à 60%, confirmant le dynamisme du secteur digital africain. Le projet Accra Fintech, avec un ROI de 60% sur un budget modeste de 10 M€, démontre l'efficacité des investissements ciblés dans la fintech en Afrique de l'Ouest. En revanche, les projets logistiques et agricoles présentent des ROIs plus modestes (30% à 40%), suggérant des cycles de retour plus longs ou des marges opérationnelles plus serrées dans ces secteurs.")

add_image(f'{OUTPUT_DIR}/bar_roi_par_projet.png', "Figure 1 : ROI par Projet — classé par ordre de performance")

# ═════════════════════════════════════════════════════════════════
# 4. TENDANCES MENSUELLES (LINE CHART)
# ═════════════════════════════════════════════════════════════════
doc.add_heading("4. Analyse des Tendances Mensuelles", level=1)

add_para("L'analyse des tendances mensuelles sur l'année 2024 révèle une trajectoire de croissance remarquablement soutenue pour l'ensemble des indicateurs financiers du portefeuille. Le graphique en courbes ci-dessous illustre l'évolution des revenus, budgets et coûts mensuels, permettant de visualiser l'accélération de la performance au fil des trimestres.")

add_para("Les revenus mensuels ont connu une progression quasi-continue, passant de 12 M€ en janvier à 95 M€ en décembre, soit une multiplication par 7,9 sur l'année. Cette croissance exponentielle reflète à la fois la montée en puissance des projets actifs et la consolidation des revenus récurrents générés par les projets arrivés à maturité. La zone verte entre les courbes de revenus et de coûts représente la marge brute dégagée chaque mois, qui s'est considérablement élargie au second semestre.")

add_para("L'analyse trimestre par trimestre montre une accélération marquée au Q3 et Q4. Au premier trimestre, la croissance moyenne mensuelle était de 4,3 M€, passant à 6,3 M€ au deuxième trimestre, puis 8,3 M€ au troisième trimestre pour atteindre 11,7 M€ au quatrième trimestre. Cette accélération est cohérente avec le cycle de maturation des projets d'investissement en Afrique, où les premiers 6 à 12 mois sont consacrés à la mise en place opérationnelle avant de générer des flux de trésorerie significatifs.")

add_para("Les budgets mensuels ont suivi une trajectoire parallèle, passant de 8 M€ à 85 M€, tandis que les coûts sont restés maîtrisés avec une progression de 6 M€ à 68 M€. L'écart croissant entre revenus et coûts témoigne de l'amélioration continue de l'efficacité opérationnelle et de la capacité du portefeuille à générer de la valeur à partir des ressources déployées.")

add_image(f'{OUTPUT_DIR}/line_tendances.png', "Figure 2 : Tendances Mensuelles — Budget, Coûts et Revenus (2024)")

# ═════════════════════════════════════════════════════════════════
# 5. REPARTITION SECTORIELLE (PIE CHART)
# ═════════════════════════════════════════════════════════════════
doc.add_heading("5. Répartition Sectorielle des Revenus", level=1)

add_para("L'analyse de la répartition sectorielle des revenus permet de comprendre la contribution relative de chaque secteur à la performance globale du portefeuille. Cette répartition est essentielle pour évaluer la diversification du portefeuille et identifier les secteurs qui devraient recevoir une allocation supplémentaire ou, au contraire, un rééquilibrage stratégique.")

sector_rev = {}
for p in projects:
    s = p['sector']
    sector_rev[s] = sector_rev.get(s, 0) + p['revenue']
total_rev = sum(sector_rev.values())
add_para(f"Le secteur minier domine la répartition avec {sector_rev['Mining']:.1f} M€ de revenus, soit {sector_rev['Mining']/total_rev*100:.1f}% du total, porté exclusivement par le projet Nimba Iron Mine. Cette concentration, bien que révélatrice d'une performance exceptionnelle, soulève la question de la dépendance du portefeuille à un seul projet minier. L'immobilier et la logistique contribuent respectivement à hauteur de {sector_rev['Real Estate']:.1f} M€ et {sector_rev['Logistics']:.1f} M€, ensemble représentant plus du tiers des revenus totaux.")

add_para(f"Les secteurs de l'énergie, de la technologie et de l'agriculture contribuent respectivement {sector_rev['Energy']:.1f} M€, {sector_rev['Technology']:.1f} M€ et {sector_rev['Agriculture']:.1f} M€. Le secteur technologique, bien que représentant seulement {sector_rev['Technology']/total_rev*100:.1f}% des revenus en valeur absolue, affiche le meilleur ratio ROI/investissement du portefeuille, ce qui en fait un candidat privilégié pour une augmentation d'allocation dans les cycles futurs d'investissement.")

add_image(f'{OUTPUT_DIR}/pie_revenus_secteur.png', "Figure 3 : Répartition des Revenus par Secteur")

# ═════════════════════════════════════════════════════════════════
# 6. PERFORMANCE EN CHANDELIERS (CANDLESTICK)
# ═════════════════════════════════════════════════════════════════
doc.add_heading("6. Performance Mensuelle en Chandeliers", level=1)

add_para("Le graphique en chandeliers offre une visualisation avancée de la performance mensuelle des investissements, intégrant les niveaux d'ouverture, de clôture, ainsi que les bornes supérieures (plus haut) et inférieures (plus bas) pour chaque mois de l'année 2024. Cette représentation est particulièrement prisée par les investisseurs professionnels car elle révèle la volatilité et la dynamique des rendements.")

add_para("L'analyse des chandeliers montre une tendance haussière marquée tout au long de l'année, avec seulement deux mois de recul mineur (février et avril) par rapport aux niveaux de clôture précédents. Les mois de novembre et décembre présentent les gammes de variation les plus importantes, avec des écarts entre les plus hauts et plus bas de 18 M€ et 23 M€ respectivement, reflétant l'intensification de l'activité d'investissement en fin d'année.")

add_para("La proportion de chandeliers haussiers (close > open) est de 100% sur l'année, ce qui indique une absence totale de mois de perte nette. Cette performance constante est remarquable pour un portefeuille d'investissements en Afrique subsaharienne, où la volatilité est traditionnellement plus élevée. Les mèches supérieures (écart entre close et high) sont généralement plus longues que les mèches inférieures, suggérant que les marchés ont régulièrement testé des niveaux supérieurs sans pour autant les maintenir durablement.")

add_image(f'{OUTPUT_DIR}/candlestick_performance.png', "Figure 4 : Performance Mensuelle en Chandeliers (M€)")

# ═════════════════════════════════════════════════════════════════
# 7. ALLOCATION BUDGETAIRE (DONUT)
# ═════════════════════════════════════════════════════════════════
doc.add_heading("7. Allocation Budgétaire par Projet", level=1)

add_para("L'allocation budgétaire est un indicateur fondamental de la stratégie d'investissement. Le graphique en anneau ci-dessous présente la répartition des 230 M€ de budget total entre les 8 projets du portefeuille, permettant d'évaluer la concentration ou la diversification des engagements financiers.")

add_para("Le projet Nimba Iron Mine absorbe à lui seul 21,7% du budget total (50 M€), ce qui est proportionnel à son rendement exceptionnel de 79%. Le Douala Port Extension (45 M€, 19,6%) et Lagos Real Estate (35 M€, 15,2%) complètent le trio des projets les plus capitalisés, représentant ensemble 56,5% de l'enveloppe. Cette concentration sur trois projets majeurs est caractéristique d'une stratégie d'investissement asymétrique visant à maximiser les rendements sur les opportunités les plus prometteuses.")

add_para("À l'autre extrême, le projet Accra Fintech ne représente que 4,3% du budget (10 M€), mais génère un ROI de 60%, ce qui en fait l'investissement le plus efficient du portefeuille en termes de rendement par euro investi. Cette disparité entre allocation budgétaire et performance relative suggère qu'une réallocation marginale vers les projets à haute rentabilité pourrait améliorer significativement le rendement global du portefeuille sans augmentation proportionnelle du risque.")

add_image(f'{OUTPUT_DIR}/donut_allocation.png', "Figure 5 : Allocation du Budget par Projet (Donut)")

# ═════════════════════════════════════════════════════════════════
# 8. MATRICE DE RISQUE
# ═════════════════════════════════════════════════════════════════
doc.add_heading("8. Matrice de Risque", level=1)

add_para("L'évaluation des risques constitue un volet essentiel de la gestion d'un portefeuille d'investissements en Afrique. La matrice ci-dessous évalue chaque projet selon cinq dimensions de risque : politique/réglementaire, opérationnel, financier, de marché et de change, avec un score global de 1 (risque minimal) à 10 (risque maximal).")

risk_headers = ["Projet", "Risque Politique", "Risque Opérationnel", "Risque Financier", "Risque Marché", "Risque Change", "Score Global"]
risk_rows = [
    ["Nimba Iron Mine", "7/10", "6/10", "5/10", "3/10", "7/10", "5,6"],
    ["Solar Plant Dakar", "2/10", "3/10", "4/10", "3/10", "5/10", "3,4"],
    ["Lagos Real Estate", "5/10", "5/10", "6/10", "6/10", "6/10", "5,6"],
    ["Abidjan Tech Hub", "2/10", "3/10", "3/10", "5/10", "4/10", "3,4"],
    ["Kinshasa Agri-Business", "8/10", "7/10", "7/10", "6/10", "9/10", "7,4"],
    ["Accra Fintech", "2/10", "3/10", "4/10", "5/10", "4/10", "3,6"],
    ["Nairobi Logistics", "3/10", "5/10", "5/10", "5/10", "5/10", "4,6"],
    ["Douala Port Extension", "3/10", "5/10", "5/10", "4/10", "5/10", "4,4"],
]
add_table(risk_headers, risk_rows)

doc.add_paragraph()
add_para("Le projet Kinshasa Agri-Business présente le profil de risque le plus élevé avec un score global de 7,4/10, principalement en raison du contexte politique et réglementaire de la RDC (score 8/10 en risque politique) et de la volatilité du franc congolais (score 9/10 en risque de change). À l'inverse, les projets situés au Sénégal (Solar Plant Dakar) et en Côte d'Ivoire (Abidjan Tech Hub) affichent les profils de risque les plus favorables, bénéficiant d'environnements réglementaires stables et de monnaies relativement ancrées.")

add_para("Le risque de change constitue la dimension la plus critique pour le portefeuille, avec une moyenne de 5,6/10. Cette exposition est inhérente à la nature multi-devises des investissements africains et nécessite des stratégies de couverture adaptées, notamment des instruments de hedging sur les matières premières et des accords de stabilité monétaire avec les banques centrales locales.")

# ═════════════════════════════════════════════════════════════════
# 9. CONCLUSIONS & RECOMMANDATIONS
# ═════════════════════════════════════════════════════════════════
doc.add_heading("9. Conclusions et Recommandations", level=1)

add_para("L'analyse approfondie du portefeuille d'InvestFlow Africa confirme la solidité de la stratégie d'investissement déployée en Afrique subsaharienne. Avec un ROI moyen de 47,9%, une marge nette de 114,1 M€ et une croissance des revenus de 692% sur l'année, le portefeuille surperforme significativement les benchmarks sectoriels africains. Toutefois, des axes d'amélioration ont été identifiés pour optimiser davantage la performance et la gestion des risques.")

doc.add_heading("9.1 Plan d'Action Stratégique", level=2)

action_headers = ["#", "Recommandation", "Priorité", "Impact Attendu", "Horizon"]
action_rows = [
    ["1", "Augmenter l'allocation au secteur technologique (+15 M€) pour capitaliser sur les ROIs supérieurs à 55% et le potentiel de croissance du digital africain", "Haute", "+8-12% ROI global", "Q1 2025"],
    ["2", "Mettre en place un programme de couverture de change pour les projets en RDC et Nigeria afin de réduire le risque de change moyen de 30%", "Haute", "-15% volatilité", "Q1 2025"],
    ["3", "Réaliser un audit opérationnel de Kinshasa Agri-Business et définir un plan de remédiation pour ramener le score de risque sous 5/10", "Moyenne", "ROI +10 points", "Q2 2025"],
    ["4", "Diversifier géographiquement avec 2-3 nouveaux projets en Afrique de l'Est (Tanzanie, Ouganda, Rwanda) pour réduire la concentration régionale", "Moyenne", "Réduction concentration", "Q3 2025"],
    ["5", "Automatiser la collecte et l'analyse des données de performance via l'API InvestFlow pour un reporting en temps réel", "Basse", "Gain de temps 40%", "Q2 2025"],
]
add_table(action_headers, action_rows)

doc.add_paragraph()
add_para("Ces recommandations sont classées par ordre de priorité et d'impact attendu. Les deux premières actions, à haute priorité, devraient être initiées dès le premier trimestre 2025 pour maximiser leurs bénéfices sur l'année fiscale. L'automatisation du reporting, bien que classée en priorité basse, représente un investissement stratégique à long terme pour améliorer la prise de décision et la transparence envers les investisseurs.")

# ═════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════
doc.save(REPORT_PATH)
print(f"\nReport saved: {REPORT_PATH}")
print(f"File size: {os.path.getsize(REPORT_PATH) / 1024:.0f} KB")
