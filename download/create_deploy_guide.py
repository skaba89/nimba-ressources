# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = '/home/z/my-project/download/'
DOCX_PATH = os.path.join(OUTPUT_DIR, 'Guide_Deploiement_Render_InvestFlow_Africa.docx')

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ── Style Configuration ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, (size, color) in enumerate([
    (Pt(22), RGBColor(0x2B, 0x6C, 0xB0)),  # Heading 1
    (Pt(16), RGBColor(0x2B, 0x6C, 0xB0)),  # Heading 2
    (Pt(13), RGBColor(0x13, 0x61, 0x2E)),  # Heading 3
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.size = size
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)

# ── Helper Functions ──
def add_paragraph(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Add shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7F7F7" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_style(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
        set_cell_shading(cell, 'F2F2F2')
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)
    return table

def add_figure(image_path, caption, width_cm=14):
    if not os.path.exists(image_path):
        add_paragraph(f'[Image non disponible: {image_path}]', italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.name = 'Arial'
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    cap_run.italic = True

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
# Add spacing before title
for _ in range(6):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guide de Deploiement Render')
run.font.name = 'Arial'
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
run.bold = True

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('InvestFlow Africa')
run.font.name = 'Arial'
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x13, 0x61, 0x2E)
run.bold = True

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_' * 60)
run.font.color.rgb = RGBColor(0xF5, 0xA5, 0x24)
run.font.size = Pt(12)

# Company
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nimba Ressources Company')
run.font.name = 'Arial'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Avril 2026')
run.font.name = 'Arial'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

# Version
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0')
run.font.name = 'Arial'
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)

# ══════════════════════════════════════════════════════════
# PAGE BREAK + HEADER/FOOTER SECTION
# ══════════════════════════════════════════════════════════
doc.add_page_break()

# Add header & footer to all subsequent sections
section = doc.sections[-1]
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = hp.add_run('Guide de Deploiement - InvestFlow Africa')
hr.font.name = 'Arial'
hr.font.size = Pt(10)
hr.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
hr.italic = True

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('Page ')
fr.font.name = 'Arial'
fr.font.size = Pt(10)
fr.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
# Page number field
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_pg = fp.add_run()
run_pg._r.append(fldChar1)
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run_pg2 = fp.add_run()
run_pg2._r.append(instrText)
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_pg3 = fp.add_run()
run_pg3._r.append(fldChar2)

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
doc.add_heading('Table des Matieres', level=1)

# Add TOC field
p_toc = doc.add_paragraph()
fldBegin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
r1 = p_toc.add_run()
r1._r.append(fldBegin)
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
r2 = p_toc.add_run()
r2._r.append(instrText)
fldSep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
r3 = p_toc.add_run()
r3._r.append(fldSep)
r4 = p_toc.add_run('[Mettre a jour la table des matieres: Ctrl+A puis F9 dans Word]')
r4.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
r4.font.italic = True
r4.font.size = Pt(10)
fldEnd = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
r5 = p_toc.add_run()
r5._r.append(fldEnd)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 1: Vue d'ensemble du projet
# ══════════════════════════════════════════════════════════
doc.add_heading('1. Vue d\'ensemble du projet', level=1)

add_paragraph(
    'InvestFlow Africa est une plateforme SaaS (Software as a Service) de pointe, concue specialement pour la gestion '
    'd\'investissements en Afrique subsaharienne. Developpee pour le compte de Nimba Ressources Company, cette application '
    'web offre une suite complete d\'outils permettant aux investisseurs, gestionnaires de fonds et decideurs economiques '
    'd\'analyser, simulator et optimiser leurs portefeuilles d\'investissement dans la region. La plateforme combine '
    'des visualisations de donnees interactives, des simulateurs de ROI en temps reel et un tableau de bord analytique '
    'puissant pour faciliter la prise de decision.'
)

add_paragraph(
    'L\'architecture technique repose sur le framework Next.js 15 avec le modele App Router, ce qui permet de beneficier '
    'd\'un rendu cote serveur optimise, d\'un routing moderne base sur le systeme de fichiers, et de performances '
    'exemplaires grace a la fonctionnalite de sortie "standalone" de Next.js. Cette fonctionnalite genere un bundle '
    'de production autonome, ideal pour les deploiements conteneurises avec Docker. L\'interface utilisateur est construite '
    'avec React 19 et React DOM, stylisee avec Tailwind CSS 4 pour un design moderne, responsive et hautement personnalisable. '
    'Le systeme de design Nimba utilise une palette de couleurs corporate associant le bleu (#2B6CB0), le vert (#13612e) et '
    'le doré (#F5A524), avec les polices Ubuntu et Plus Jakarta Sans pour une identite visuelle forte et professionnelle.'
)

add_paragraph(
    'La plateforme integre egalement Recharts pour la generation de 10 types de graphiques interactifs (barres, lignes, '
    'camemberts, radar, donut, aires empilees, chandeliers, entonnoirs, cascades et barres empilees), ainsi que Lucide React '
    'pour une iconographie coherente et moderne. Le traitement de fichiers CSV est assure par PapaParse, permettant aux '
    'utilisateurs d\'importer leurs propres donnees d\'investissement directement dans la plateforme. L\'ensemble de ces '
    'technologies est conteneurise dans une image Docker multi-stage optimisee, garantissant un deploiement rapide, '
    'reproductible et scalable sur n\'importe quelle infrastructure cloud.'
)

# Tech stack table
doc.add_heading('Stack technique', level=2)
add_table_with_style(
    ['Categorie', 'Technologie', 'Version', 'Role'],
    [
        ['Framework', 'Next.js', '15 (App Router)', 'Framework React full-stack'],
        ['Frontend', 'React / React DOM', '19', 'Interface utilisateur'],
        ['Styling', 'Tailwind CSS', '4', 'Utilitaires CSS'],
        ['Graphiques', 'Recharts', '3.8', 'Visualisations interactives'],
        ['Icones', 'Lucide React', '1.7', 'Systeme iconographique'],
        ['Donnees', 'PapaParse', '5.5', 'Import/export CSV'],
        ['Conteneur', 'Docker', 'Multi-stage', 'Deploiement production'],
        ['Hebergement', 'Render.com', 'Free/Paid', 'Infrastructure cloud'],
    ],
    col_widths=[3.5, 3.5, 2.5, 5.5]
)

# ══════════════════════════════════════════════════════════
# SECTION 2: Prerequis
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Prerequis', level=1)

add_paragraph(
    'Avant de proceder au deploiement d\'InvestFlow Africa sur Render, plusieurs prerequis doivent etre satisfaits '
    'pour garantir une installation fluide et sans erreur. Ce guide suppose que vous disposez d\'acces au depot '
    'GitHub du projet et que vous avez les permissions necessaires pour creer des services sur la plateforme Render. '
    'L\'ensemble du processus est concu pour etre realise entierement depuis un navigateur web, sans necessite '
    'd\'installation locale de Node.js, npm ou tout autre outil de developpement.'
)

doc.add_heading('2.1 Compte et acces', level=2)
add_bullet('Un compte GitHub actif avec acces en lecture/ecriture au depot https://github.com/skaba89/nimba-ressources')
add_bullet('Un compte Render.com (inscription gratuite sur https://render.com). Le plan Free est suffisant pour un deploiement initial et des tests.')
add_bullet('Une connexion internet stable pour le telechargement des images Docker et le build de l\'application.')

doc.add_heading('2.2 Connaissances requises', level=2)
add_paragraph(
    'Aucune connaissance approfondie en DevOps n\'est requise pour suivre ce guide. Cependant, une comprehension '
    'de base des concepts suivants sera utile pour le deroulement optimal des operations : la notion de conteneur '
    'Docker (bien que le build soit automatise), le fonctionnement de Git et des depots distants, ainsi que les '
    'principes generaux de deploiement d\'applications web. Chaque etape est detaillee avec des captures d\'ecran '
    'et des instructions precisent pour guider les utilisateurs de tous niveaux de competences techniques.'
)

doc.add_heading('2.3 Limites du plan gratuit Render', level=2)
add_paragraph(
    'Le plan gratuit de Render.com impose certaines limitations qu\'il convient de connaitre a l\'avance : '
    'le service est mis en veille apres 15 minutes d\'inactivite (cold start d\'environ 30-60 secondes au prochain acces), '
    'la RAM est limitee a 512 Mo, le disque a 1 Go, et aucun domaine personnalise n\'est inclus. Pour un usage '
    'en production avec un trafic regulier, il est recommande de passer au plan Starter a 7$/mois qui supprime '
    'ces restrictions et offre une meilleure performances.'
)

# ══════════════════════════════════════════════════════════
# SECTION 3: Configuration Docker
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Configuration Docker', level=1)

add_paragraph(
    'Le deploiement d\'InvestFlow Africa repose sur une architecture Docker multi-stage optimisee, specialement '
    'concue pour minimiser la taille de l\'image finale tout en garantissant un build reproductible et securise. '
    'Le Dockerfile est structure en trois etapes distinctes : la premiere installe les dependances, la seconde '
    'compile l\'application Next.js, et la troisieme produit une image de production ultra-legere ne contenant que '
    'les fichiers strictement necessaires a l\'execution. Cette approche permet de reduire la taille de l\'image '
    'de plus de 74% par rapport a un build monolithique traditionnel, passant de 318 Mo (builder) a seulement '
    '82 Mo (image finale).'
)

doc.add_heading('3.1 Architecture Multi-Stage', level=2)
add_paragraph(
    'L\'architecture multi-stage du Dockerfile est un pattern de bonnes pratiques en conteneurisation qui consiste '
    'a diviser le processus de build en plusieurs etapes independantes. Chaque etape utilise une image de base '
    'distincte (node:20-alpine) et ne conserve que les artefacts necessaires a l\'etape suivante. Cela permet '
    'd\'eliminer les outils de compilation, les fichiers source intermediaires et les dependances de developpement '
    'de l\'image finale, resulting en un conteneur de production plus petit, plus securise et plus rapide a deployer.'
)

add_figure(
    os.path.join(OUTPUT_DIR, 'chart_docker_architecture.png'),
    'Figure 3-1 : Architecture Docker Multi-Stage - Reduction de la taille de l\'image',
    width_cm=14
)

doc.add_heading('3.2 Explication du Dockerfile', level=2)
add_paragraph(
    'Le Dockerfile commence par l\'etape "deps" qui installe uniquement les dependances de production avec '
    'npm ci --only=production. L\'option --ci garantit une installation deterministe basee sur le fichier '
    'package-lock.json, tandis que --only=production exclut les devDependencies (eslint, typescript, etc.) '
    'qui ne sont pas necessaires en production. L\'etape "builder" copie les node_modules de production, '
    'reinstalle toutes les dependances (incluant les devDependencies pour le build), compile l\'application '
    'avec npm run build, et genere le bundle standalone de Next.js dans le repertoire .next/standalone/. '
    'Enfin, l\'etape "runner" ne copie que le bundle standalone, les assets statiques (.next/static) et '
    'les fichiers publics (images, logo). Aucun node_modules complet n\'est inclus dans l\'image finale.'
)

doc.add_heading('3.3 Configuration du port', level=2)
add_paragraph(
    'Le port d\'ecoute est configure sur 10000, qui est le port par defaut utilise par Render.com pour les services web. '
    'Render injecte automatiquement la variable d\'environnement PORT lors du deploiement, et Next.js standalone '
    'detecte cette variable pour demarrer le serveur HTTP sur le bon port. Le HOSTNAME est defini sur "0.0.0.0" '
    'pour ecouter sur toutes les interfaces reseau, ce qui est indispensable dans un environnement conteneurise '
    'ou l\'application doit etre accessible depuis l\'exterieur du conteneur. Le health check est configure avec '
    'wget pour verifier la disponibilite de l\'application toutes les 30 secondes, avec un delai de demarrage '
    'de 40 secondes pour laisser le temps a Next.js de se charger.'
)

add_code_block('# Extrait du Dockerfile - Configuration du port\nENV PORT=10000\nENV HOSTNAME="0.0.0.0"\nEXPOSE 10000\nHEALTHCHECK --interval=30s --timeout=10s --start-period=40s --retries=3 \\\n  CMD wget --no-verbose --tries=1 --spider http://localhost:10000/ || exit 1')

# ══════════════════════════════════════════════════════════
# SECTION 4: Deploiement Render - Methode Dashboard
# ══════════════════════════════════════════════════════════
doc.add_heading('4. Deploiement sur Render - Methode Dashboard', level=1)

add_paragraph(
    'La methode Dashboard est la facon la plus intuitive et visuelle de deployer InvestFlow Africa sur Render.com. '
    'Elle consiste a utiliser l\'interface graphique de Render pour connecter votre depot GitHub, configurer les '
    'parametres de build, et lancer le deploiement en quelques clics. Cette methode est particulierement recommandee '
    'pour les premiers deploiements car elle offre une visibilite complete sur chaque etape du processus et permet '
    'de detecter rapidement toute erreur de configuration. Le deploiement complet prend environ 2 a 3 minutes '
    'lors de la premiere exécution, puis environ 1 minute pour les deploiements suivants (auto-deploy sur git push).'
)

doc.add_heading('Etape 1 : Connexion a Render.com', level=2)
add_paragraph(
    'Rendez-vous sur https://dashboard.render.com et connectez-vous avec votre compte GitHub, GitLab ou Google. '
    'Si vous n\'avez pas encore de compte, cliquez sur "Get Started for Free" et inscrivez-vous en quelques secondes. '
    'Une fois connecte, vous arrivez sur le tableau de bord principal qui affiche la liste de vos services actifs. '
    'Si c\'est votre premier depot, cette liste sera vide.'
)

doc.add_heading('Etape 2 : Creation d\'un nouveau Web Service', level=2)
add_paragraph(
    'Cliquez sur le bouton "New" en haut a droite du tableau de bord, puis selectionnez "Web Service" dans le menu '
    'deroulant. Render vous demande alors de connecter un depot Git. Cliquez sur "Connect a repository", recherchez '
    '"nimba-ressources" dans la liste de vos depots GitHub, et selectionnez-le. Si le depot n\'apparaît pas, '
    'verifiez que vous avez bien les droits d\'acces et que le depot est public ou que Render a les permissions OAuth.'
)

doc.add_heading('Etape 3 : Configuration du service', level=2)
add_paragraph(
    'Une fois le depot connecte, vous accedez a la page de configuration du service. Les champs a remplir sont les suivants :'
)

add_table_with_style(
    ['Parametre', 'Valeur', 'Description'],
    [
        ['Name', 'investflow-africa', 'Nom du service (devient le sous-domaine)'],
        ['Region', 'Frankfurt (eu-west)', 'Region la plus proche de l\'Afrique de l\'Ouest'],
        ['Runtime', 'Docker', 'Render detecte automatiquement le Dockerfile'],
        ['Plan', 'Free / Starter', 'Free pour tester, Starter (7$/mois) pour la prod'],
        ['Branch', 'main', 'Branche GitHub a deployer'],
    ],
    col_widths=[3, 4, 9]
)

doc.add_heading('Etape 4 : Variables d\'environnement', level=2)
add_paragraph(
    'Dans la section "Environment Variables", ajoutez les variables suivantes si elles ne sont pas detectees '
    'automatiquement. Render injecte la variable PORT par defaut, mais il est recommande de la definir explicitement '
    'pour des raisons de clarte et de compatibilite. La variable NEXT_TELEMETRY_DISABLED est mise a "1" pour desactiver '
    'la telemetrie Next.js en production et optimiser les performances.'
)

add_code_block('NODE_ENV=production\nNEXT_TELEMETRY_DISABLED=1\nPORT=10000')

doc.add_heading('Etape 5 : Lancement du deploiement', level=2)
add_paragraph(
    'Cliquez sur le bouton "Create Web Service" en bas de la page. Render lance immediatement le processus de build : '
    'il clone le depot, construit l\'image Docker selon le Dockerfile multi-stage (environ 90-120 secondes), '
    'puis demarre le conteneur de production. Vous pouvez suivre l\'avancement en temps reel dans les logs '
    'de deployment accessibles via l\'onglet "Logs" de votre service. Une fois le deploiement termine, '
    'Render affiche un indicateur vert "Live" a cote du nom de votre service.'
)

doc.add_heading('Etape 6 : Verification', level=2)
add_paragraph(
    'Votre application est desormais accessible a l\'adresse https://investflow-africa.onrender.com. '
    'Cliquez sur l\'URL fournie dans le tableau de bord Render pour ouvrir la plateforme dans votre navigateur. '
    'Verifiez que toutes les sections s\'affichent correctement : le hero avec l\'image de fond, le tableau de bord '
    'avec les KPI, le simulateur ROI interactif, les 10 graphiques Recharts, la galerie de projets, les temoignages, '
    'les tarifs et le footer. Si vous constatez une erreur 502, attendez quelques instants (le health check peut '
    'prendre jusqu\'a 40 secondes) et rafraichissez la page.'
)

# ══════════════════════════════════════════════════════════
# SECTION 5: Deploiement Render - Methode render.yaml
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Deploiement via Blueprint (render.yaml)', level=1)

add_paragraph(
    'Render Blueprint est une fonctionnalite qui permet de definir l\'ensemble de votre infrastructure comme code (IaC) '
    'dans un fichier render.yaml a la racine de votre depot. Cette methode offre plusieurs avantages : elle garantit '
    'la reproductibilite des deploiements, facilite la collaboration en equipe, permet le versioning de la configuration '
    'd\'infrastructure avec le code source, et automatise la creation de services sans intervention manuelle sur le dashboard. '
    'Le fichier render.yaml du projet InvestFlow Africa est deja configure et pret a l\'emploi.'
)

doc.add_heading('5.1 Contenu du fichier render.yaml', level=2)
add_code_block('''services:
  - type: web
    name: investflow-africa
    runtime: docker
    plan: free
    region: frankfurt
    envVars:
      - key: NODE_ENV
        value: production
      - key: NEXT_TELEMETRY_DISABLED
        value: "1"
    healthCheckPath: /''')

add_paragraph(
    'Ce fichier YAML definit un service web de type Docker, nomme "investflow-africa", deploye dans la region de '
    'Frankfurt (la plus proche de l\'Afrique de l\'Ouest pour une latence optimale). Le plan est configure sur "free" '
    'par defaut, mais peut etre modifie en "starter" ou "standard" directement dans le fichier ou via le dashboard. '
    'Les variables d\'environnement sont definies inline, et le health check est configure sur le chemin "/" pour '
    'verifier la disponibilite de la page d\'accueil.'
)

doc.add_heading('5.2 Procedure de deploiement', level=2)
add_bullet('Assurez-vous que le fichier render.yaml est present a la racine du depot et commis sur la branche main.')
add_bullet('Connectez-vous a Render.com et accedez au tableau de bord.')
add_bullet('Cliquez sur "New" puis "Blueprint" au lieu de "Web Service".')
add_bullet('Selectionnez le depot nimba-ressources et la branche main.')
add_bullet('Render analyse automatiquement le fichier render.yaml et propose la creation du service.')
add_bullet('Cliquez sur "Apply" pour lancer le deploiement.')
add_paragraph(
    'L\'avantage principal de cette methode est que toute modification du fichier render.yaml (changement de region, '
    'de plan, ajout de variables d\'environnement) est prise en compte automatiquement lors du prochain push sur '
    'la branche main, sans intervention manuelle sur le dashboard de Render.'
)

# ══════════════════════════════════════════════════════════
# SECTION 6: Deploiement local avec Docker
# ══════════════════════════════════════════════════════════
doc.add_heading('6. Deploiement local avec Docker', level=1)

add_paragraph(
    'Pour tester localement l\'application avant le deploiement sur Render, vous pouvez utiliser Docker Desktop '
    'ou Docker Engine directement sur votre machine. Le deploiement local reproduit exactement les memes conditions '
    'que l\'environnement Render, garantissant que ce qui fonctionne en local fonctionnera egalement en production. '
    'Cette etape est fortement recommandee pour verifier le bon fonctionnement de l\'application et diagnostiquer '
    'd\'eventuels problemes avant de pousser le code en production.'
)

doc.add_heading('6.1 Build de l\'image Docker', level=2)
add_code_block('# Construire l\'image Docker (depuis la racine du projet)\ndocker build -t investflow-africa:latest .')
add_paragraph(
    'Le build prend environ 2 minutes lors de la premiere execution (telechargement des images de base + installation '
    'des dependances + compilation Next.js). Les builds suivants sont significativement plus rapides grace au cache '
    'Docker, qui ne re-execute que les etapes ayant change depuis le dernier build. L\'image finale pese environ '
    '82 Mo, ce qui est tres compact pour une application Next.js complete avec toutes ses dependances.'
)

doc.add_heading('6.2 Lancement avec docker run', level=2)
add_code_block('# Lancer le conteneur\ndocker run -d -p 10000:10000 --name investflow investflow-africa:latest\n\n# Acceder a l\'application\n# Ouvrir http://localhost:10000 dans le navigateur')
add_paragraph(
    'L\'option -d lance le conteneur en mode detache (arriere-plan), -p 10000:10000 mappe le port 10000 du conteneur '
    'vers le port 10000 de votre machine hote, et --name investflow donne un nom lisible au conteneur pour faciliter '
    'la gestion ulterieure (logs, stop, rm). Une fois le conteneur demarre, ouvrez votre navigateur a l\'adresse '
    'http://localhost:10000 pour acceder a la plateforme InvestFlow Africa.'
)

doc.add_heading('6.3 Lancement avec docker-compose', level=2)
add_code_block('# Lancer avec docker-compose\ndocker-compose up -d\n\n# Voir les logs\ndocker-compose logs -f\n\n# Arreter le service\ndocker-compose down')
add_paragraph(
    'Docker Compose simplifie la gestion des conteneurs en lisant la configuration depuis le fichier docker-compose.yml. '
    'La commande up -d lance le service en arriere-plan, logs -f affiche les logs en temps reel (suivi continu), '
    'et down arrete et supprime le conteneur. Le fichier docker-compose.yml du projet inclut egalement une '
    'configuration de healthcheck qui surveille automatiquement la disponibilite de l\'application.'
)

# ══════════════════════════════════════════════════════════
# SECTION 7: Domaine personnalise
# ══════════════════════════════════════════════════════════
doc.add_heading('7. Configuration du domaine personnalise', level=1)

add_paragraph(
    'Render.com permet d\'associer un domaine personnalise (par exemple, app.nimba-ressources.com) a votre service '
    'web. Cette fonctionnalite est disponible sur les plans Starter et superieurs (pas sur le plan Free). Render '
    'gere automatiquement la configuration SSL/TLS via Let\'s Encrypt, ce qui signifie que votre application sera '
    'accessible en HTTPS sans aucune configuration supplementaire de certificat. Le processus de configuration '
    'est simple et se fait en quelques etapes depuis le dashboard Render.'
)

doc.add_heading('7.1 Etapes de configuration', level=2)
add_bullet('Accedez a votre service investflow-africa sur le dashboard Render.')
add_bullet('Cliquez sur "Settings" puis sur "Custom Domain".')
add_bullet('Entrez votre domaine personnalise (ex: investflow.nimba-ressources.com).')
add_bullet('Render affiche les enregistrements DNS a configurer chez votre registrar.')
add_bullet('Ajoutez un enregistrement CNAME pointant vers le domaine Render (xxxxxxxx.onrender.com).')
add_bullet('Attendez la propagation DNS (de quelques minutes a 48 heures maximum).')
add_bullet('Render emet automatiquement un certificat SSL/TLS une fois le domaine verifie.')

add_paragraph(
    'Pour Nimba Ressources Company, il est recommande d\'utiliser un sous-domaine specifique comme investflow.nimba-ressources.com '
    'ou app.nimba-ressources.com pour separer l\'application des autres services web de l\'entreprise. Si vous possedez '
    'deja un site vitrine sur le domaine principal, un sous-domaine permet de coexister sans conflit.'
)

# ══════════════════════════════════════════════════════════
# SECTION 8: Deploiement comparatif
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Analyse comparative des temps de deploiement', level=1)

add_paragraph(
    'Le choix de la strategie de deploiement a un impact direct sur la rapidite de mise a jour et la reactivite '
    'de l\'equipe face aux evolutions du produit. L\'architecture Docker multi-stage d\'InvestFlow Africa permet '
    'de beneficier de temps de deploiement competitifs par rapport aux solutions serverless comme Vercel, tout en '
    'conservant la flexibilite d\'un conteneur standard qui peut etre deploye sur n\'importe quelle plateforme cloud.'
)

add_figure(
    os.path.join(OUTPUT_DIR, 'chart_deploy_time.png'),
    'Figure 8-1 : Temps de deploiement comparatif selon la methode utilisee',
    width_cm=13
)

add_paragraph(
    'Le premier deploiement sur Render prend environ 180 secondes en raison de la construction initiale de l\'image '
    'Docker (telechargement des images de base, installation des dependances npm, compilation Next.js). Cependant, '
    'les deploiements suivants sont significativement plus rapides (environ 60 secondes) grace au cache Docker qui '
    'ne re-compile que les fichiers modifies. Le deploiement local avec Docker prend environ 120 secondes pour '
    'un build complet. Ces temps restent tout a fait acceptables pour une application de cette envergure et '
    'permettent une iteration rapide sur les fonctionnalites.'
)

# ══════════════════════════════════════════════════════════
# SECTION 9: Depannage (Troubleshooting)
# ══════════════════════════════════════════════════════════
doc.add_heading('9. Depannage (Troubleshooting)', level=1)

add_paragraph(
    'Cette section recense les problemes les plus frequemment rencontres lors du deploiement d\'InvestFlow Africa '
    'sur Render et fournit des solutions detaillees pour chacun d\'entre eux. Les erreurs de deploiement sont '
    'normales et font partie du cycle de developpement. L\'important est de savoir les identifier rapidement '
    'grace aux logs et d\'appliquer la solution appropriee. Render fournit des logs en temps reel dans le '
    'dashboard qui sont essentiels pour le diagnostic.'
)

doc.add_heading('9.1 Tableau des erreurs courantes', level=2)
add_table_with_style(
    ['Erreur', 'Cause probable', 'Solution'],
    [
        ['Build failed', 'Dockerfile invalide ou dependance manquante', 'Verifier le Dockerfile, executer docker build localement pour tester'],
        ['Port binding error', 'Application ecoute sur le mauvais port', 'S\'assurer que PORT=10000 est defini, utiliser process.env.PORT dans le code'],
        ['Health check failed', 'Delai de demarrage insuffisant', 'Augmenter start-period a 60s dans le Dockerfile'],
        ['Image too large', 'Fichiers inutiles inclus dans le build', 'Verifier le fichier .dockerignore, exclure node_modules, .next, download/'],
        ['502 Bad Gateway', 'Conteneur crash au demarrage', 'Consulter les logs Render, verifier les variables d\'environnement'],
        ['Cold start lent', 'Plan gratuit, service en veille', 'Passer au plan Starter ou effectuer un ping regulier'],
        ['CSS non charge', 'Assets statiques manquants', 'Verifier que .next/static et public/ sont copies dans le Dockerfile'],
        ['Out of memory', '512 Mo insuffisant pour le build', 'Augmenter la RAM dans les parametres Render (plan superieur)'],
    ],
    col_widths=[3.5, 4.5, 8]
)

doc.add_heading('9.2 Commandes de diagnostic', level=2)
add_code_block('# Build local pour tester le Dockerfile\ndocker build -t investflow-test . 2>&1 | tee build.log\n\n# Lancer le conteneur et voir les logs\ndocker run --rm -p 10000:10000 investflow-test\n\n# Verifier les logs Render\n# Dashboard > investflow-africa > Logs > Live\n\n# Test de sante manuel\ncurl -I https://investflow-africa.onrender.com/')

doc.add_heading('9.3 Bonnes pratiques de debug', level=2)
add_paragraph(
    'En cas de probleme, la premiere chose a faire est de reproduire l\'erreur localement avec Docker. '
    'Si le build echoue localement, le probleme vient du code ou du Dockerfile. Si le build reussit localement '
    'mais echoue sur Render, verifiez les differences d\'environnement (version de Node, variables d\'environnement, '
    'fichiers presents dans le depot). Il est egalement recommande de ne pas pousser de fichiers generes (images, '
    'rapports, captures) dans le depot GitHub car ils alourdissent inutilement le build et peuvent causer des '
    'erreurs de depassement de taille. Le fichier .gitignore du projet est configure pour exclure ces fichiers.'
)

# ══════════════════════════════════════════════════════════
# SECTION 10: Bonnes pratiques
# ══════════════════════════════════════════════════════════
doc.add_heading('10. Bonnes pratiques et recommandations', level=1)

doc.add_heading('10.1 Gestion des variables d\'environnement', level=2)
add_paragraph(
    'Les variables d\'environnement sensibles (cles API, identifiants de base de donnees, secrets) ne doivent '
    'jamais etre commises dans le depot Git. Elles doivent etre configurees exclusivement via le dashboard Render, '
    'dans la section "Environment Variables" du service. Render chiffre automatiquement ces variables et les injecte '
    'au moment du build et de l\'execution. Pour le developpement local, utilisez un fichier .env.local (deja '
    'exclu du depot par le .gitignore) qui ne sera jamais versionne ni inclus dans l\'image Docker.'
)

doc.add_heading('10.2 Auto-deploy et CI/CD', level=2)
add_paragraph(
    'Render propose un auto-deploy automatique qui surveille votre depot GitHub et declenche un nouveau deploiement '
    'a chaque push sur la branche configuree (main par defaut). Cette fonctionnalite est activee par defaut et '
    'constitue un pipeline CI/CD simple et efficace sans configuration supplementaire. Pour les workflows plus '
    'complexes, Render supporte egalement les deploiements manuels via le dashboard, les deploiements conditionnels '
    'bases sur des webhooks, et l\'integration avec des services CI externes comme GitHub Actions.'
)

doc.add_heading('10.3 Monitoring et alertes', level=2)
add_paragraph(
    'Le dashboard Render fournit des metriques en temps reel incluant le CPU, la memoire, le reseau et les '
    'temps de reponse HTTP. Configurez des alertes email pour etre notifie en cas de depassement de seuils '
    '(CPU > 80%, memoire > 90%, erreurs HTTP > 5%). Pour un monitoring avance, vous pouvez integrer des '
    'outils externes comme Sentry pour le suivi des erreurs JavaScript, ou UptimeRobot pour la surveillance '
    'de la disponibilite avec des alertes Slack ou email.'
)

doc.add_heading('10.4 Scalabilite', level=2)
add_paragraph(
    'La conception d\'InvestFlow Africa est par defaut stateless (sans etat), ce qui signifie qu\'elle peut '
    'etre mise a l\'echelle horizontalement en creant plusieurs instances du service sur Render. Chaque instance '
    'est identique et peut traiter les requetes independamment. Pour activer le scaling, accedez aux parametres '
    'du service et configurez le nombre minimum et maximum d\'instances. Render gere automatiquement la '
    'repartition de charge (load balancing) entre les instances via son reverse proxy interne.'
)

doc.add_heading('10.5 Sauvegarde et rollback', level=2)
add_paragraph(
    'Chaque deploiement sur Render cree un snapshot de l\'image Docker qui peut etre reinstalle en un clic '
    'en cas de probleme. Pour effectuer un rollback, accedez a la section "Events" de votre service, '
    'identifiez le dernier deploiement fonctionnel, et cliquez sur "Deploy" pour le reactiver. En parallelle, '
    'il est recommande de tagger les versions stables dans Git avec des tags semantiques (v1.0.0, v1.1.0) '
    'pour faciliter le suivi des versions deployees et la gestion des releases.'
)

# ── Save Document ──
doc.save(DOCX_PATH)
print(f'Document saved to: {DOCX_PATH}')
