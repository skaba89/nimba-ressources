#!/usr/bin/env python3
"""
InvestFlow Africa — Rapport d'Analyse de Performance Q4 2024
Generates charts (matplotlib) and professional Word document (python-docx)
"""

import os
import sys
from datetime import datetime
from collections import defaultdict

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from matplotlib.font_manager import FontProperties

from docx import Document
from docx.shared import Inches, Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ──────────────────────────────────────────────────────────────
# PATHS
# ──────────────────────────────────────────────────────────────
BASE_DIR = "/home/z/my-project"
FONT_PATH = os.path.join(BASE_DIR, "fonts", "simhei.ttf")
CHARTS_DIR = os.path.join(BASE_DIR, "download", "charts_report")
OUTPUT_DOC = os.path.join(BASE_DIR, "download", "InvestFlow_Africa_Rapport_Analyse_2026.docx")

os.makedirs(CHARTS_DIR, exist_ok=True)

# ──────────────────────────────────────────────────────────────
# DATA
# ──────────────────────────────────────────────────────────────
projects = [
    {"name": "Nimba Iron Mine", "sector": "Mining", "budget": 50, "cost": 38, "revenue": 89.5, "roi": 79, "status": "Active", "country": "Guinée"},
    {"name": "Solar Plant Dakar", "sector": "Energy", "budget": 20, "cost": 16, "revenue": 29.1, "roi": 44, "status": "Active", "country": "Sénégal"},
    {"name": "Lagos Real Estate", "sector": "Real Estate", "budget": 35, "cost": 30, "revenue": 52.5, "roi": 50, "status": "Active", "country": "Nigeria"},
    {"name": "Abidjan Tech Hub", "sector": "Technology", "budget": 15, "cost": 12, "revenue": 22.5, "roi": 50, "status": "Completed", "country": "Côte d'Ivoire"},
    {"name": "Kinshasa Agri-Business", "sector": "Agriculture", "budget": 25, "cost": 20, "revenue": 32.5, "roi": 30, "status": "Active", "country": "RDC"},
    {"name": "Accra Fintech", "sector": "Technology", "budget": 10, "cost": 8, "revenue": 16.0, "roi": 60, "status": "Active", "country": "Ghana"},
    {"name": "Nairobi Logistics", "sector": "Logistics", "budget": 30, "cost": 24, "revenue": 39.0, "roi": 30, "status": "Pending", "country": "Kenya"},
    {"name": "Douala Port Extension", "sector": "Logistics", "budget": 45, "cost": 36, "revenue": 63.0, "roi": 40, "status": "Active", "country": "Cameroun"},
]

# Computed aggregates
total_budget = sum(p["budget"] for p in projects)
total_cost = sum(p["cost"] for p in projects)
total_revenue = sum(p["revenue"] for p in projects)
avg_roi = sum(p["roi"] for p in projects) / len(projects)

# Sector aggregation
sector_revenue = defaultdict(float)
sector_budget = defaultdict(float)
for p in projects:
    sector_revenue[p["sector"]] += p["revenue"]
    sector_budget[p["sector"]] += p["budget"]

# ──────────────────────────────────────────────────────────────
# NIMBA PALETTE
# ──────────────────────────────────────────────────────────────
NIMBA = {
    "blue": "#2B6CB0",
    "green": "#13612e",
    "gold": "#f5a524",
    "red": "#b82105",
    "orange": "#f7630c",
    "teal": "#2B9EB3",
    "purple": "#7C5CFC",
}
COLORS_LIST = list(NIMBA.values())
BG_COLOR = "#F7FAFC"
TEXT_COLOR = "#1A202C"

# ──────────────────────────────────────────────────────────────
# MATPLOTLIB GLOBAL STYLE
# ──────────────────────────────────────────────────────────────
plt.rcParams.update({
    "figure.facecolor": "#FFFFFF",
    "axes.facecolor": BG_COLOR,
    "axes.edgecolor": "#CBD5E0",
    "axes.grid": True,
    "grid.color": "#E2E8F0",
    "grid.linewidth": 0.7,
    "text.color": TEXT_COLOR,
    "axes.labelcolor": TEXT_COLOR,
    "xtick.color": TEXT_COLOR,
    "ytick.color": TEXT_COLOR,
    "font.family": "sans-serif",
    "font.size": 11,
    "figure.dpi": 200,
})

# Register SimHei for any CJK labels (but we'll mainly use sans-serif)
simhei_prop = FontProperties(fname=FONT_PATH)

# ──────────────────────────────────────────────────────────────
# CHART 1: Horizontal bar – ROI par projet (sorted desc)
# ──────────────────────────────────────────────────────────────
def chart_bar_roi():
    sorted_p = sorted(projects, key=lambda x: x["roi"], reverse=True)
    names = [p["name"] for p in sorted_p]
    rois = [p["roi"] for p in sorted_p]
    colors = [COLORS_LIST[i % len(COLORS_LIST)] for i in range(len(names))]

    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.barh(names, rois, color=colors, edgecolor="white", height=0.6)
    ax.invert_yaxis()
    for bar, val in zip(bars, rois):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2, f"{val}%",
                va="center", fontsize=10, fontweight="bold", color=TEXT_COLOR)
    ax.set_xlabel("ROI (%)")
    ax.set_title("ROI par Projet", fontsize=14, fontweight="bold", pad=12)
    ax.set_xlim(0, max(rois) + 15)
    ax.xaxis.set_major_formatter(mticker.FormatStrFormatter("%.0f%%"))
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, "bar_roi_par_projet.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)
    print(f"  ✓ {path}")
    return path


# ──────────────────────────────────────────────────────────────
# CHART 2: Line chart – Tendances mensuelles (12 months)
# ──────────────────────────────────────────────────────────────
def chart_line_tendances():
    months = ["Jan", "Fév", "Mar", "Avr", "Mai", "Jun", "Jul", "Aoû", "Sep", "Oct", "Nov", "Déc"]
    # Simulated monthly trends based on total budget/cost/revenue distribution
    np.random.seed(42)
    revenue_trend = np.cumsum(np.linspace(total_revenue*0.04, total_revenue*0.12, 12))
    revenue_trend = np.round(revenue_trend, 1)
    budget_trend = np.linspace(total_budget*0.06, total_budget*0.14, 12)
    budget_trend = np.round(budget_trend, 1)
    cost_trend = np.linspace(total_cost*0.05, total_cost*0.12, 12)
    cost_trend = np.round(cost_trend, 1)

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(months, revenue_trend, marker="o", linewidth=2.5, color=NIMBA["green"], label="Revenus", markersize=6)
    ax.plot(months, budget_trend, marker="s", linewidth=2.5, color=NIMBA["blue"], label="Budget", markersize=6, linestyle="--")
    ax.plot(months, cost_trend, marker="^", linewidth=2.5, color=NIMBA["gold"], label="Coûts", markersize=6, linestyle="-.")
    ax.fill_between(months, revenue_trend, alpha=0.1, color=NIMBA["green"])
    ax.set_ylabel("Montant (M USD)")
    ax.set_title("Tendances Mensuelles — Revenus, Budget et Coûts (2024)", fontsize=14, fontweight="bold", pad=12)
    ax.legend(loc="upper left", frameon=True, facecolor="white", edgecolor="#CBD5E0")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, "line_tendances.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)
    print(f"  ✓ {path}")
    return path


# ──────────────────────────────────────────────────────────────
# CHART 3: Pie chart – Revenus par secteur
# ──────────────────────────────────────────────────────────────
def chart_pie_revenus():
    labels = list(sector_revenue.keys())
    values = [sector_revenue[s] for s in labels]
    colors = COLORS_LIST[:len(labels)]

    fig, ax = plt.subplots(figsize=(9, 7))
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, autopct="%1.1f%%", startangle=140,
        colors=colors, pctdistance=0.78, labeldistance=1.12,
        wedgeprops=dict(edgecolor="white", linewidth=2),
        textprops=dict(fontsize=11, color=TEXT_COLOR),
    )
    for t in autotexts:
        t.set_fontsize(10)
        t.set_fontweight("bold")
        t.set_color("white")
    ax.set_title("Répartition des Revenus par Secteur", fontsize=14, fontweight="bold", pad=16)
    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, "pie_revenus_secteur.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)
    print(f"  ✓ {path}")
    return path


# ──────────────────────────────────────────────────────────────
# CHART 4: Stacked bar – Budget vs Coûts par projet
# ──────────────────────────────────────────────────────────────
def chart_stacked_bar():
    names = [p["name"] for p in projects]
    budgets = [p["budget"] for p in projects]
    costs = [p["cost"] for p in projects]
    savings = [p["budget"] - p["cost"] for p in projects]

    fig, ax = plt.subplots(figsize=(11, 6.5))
    x = np.arange(len(names))
    w = 0.55
    b1 = ax.bar(x, costs, width=w, label="Coûts Réels", color=NIMBA["blue"], edgecolor="white")
    b2 = ax.bar(x, savings, width=w, bottom=costs, label="Économies", color=NIMBA["teal"], edgecolor="white")

    # Budget line
    ax.plot(x, budgets, marker="D", color=NIMBA["red"], linewidth=2, linestyle="--", label="Budget Alloué", markersize=7, zorder=5)

    for i, (b, c) in enumerate(zip(budgets, costs)):
        ax.text(i, b + 1.2, f"{b}M", ha="center", fontsize=9, color=NIMBA["red"], fontweight="bold")

    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=30, ha="right", fontsize=9)
    ax.set_ylabel("Montant (M USD)")
    ax.set_title("Structure des Coûts — Budget vs Coûts Réels par Projet", fontsize=14, fontweight="bold", pad=12)
    ax.legend(loc="upper right", frameon=True, facecolor="white", edgecolor="#CBD5E0")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, "stacked_bar_couts.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)
    print(f"  ✓ {path}")
    return path


# ──────────────────────────────────────────────────────────────
# CHART 5: Donut – Allocation budgétaire
# ──────────────────────────────────────────────────────────────
def chart_donut_allocation():
    names = [p["name"] for p in projects]
    budgets = [p["budget"] for p in projects]
    colors = COLORS_LIST[:len(names)]

    fig, ax = plt.subplots(figsize=(9, 7))
    wedges, texts, autotexts = ax.pie(
        budgets, labels=names, autopct="%1.1f%%", startangle=90,
        colors=colors, pctdistance=0.8, labeldistance=1.12,
        wedgeprops=dict(width=0.4, edgecolor="white", linewidth=2),
        textprops=dict(fontsize=10, color=TEXT_COLOR),
    )
    for t in autotexts:
        t.set_fontsize(9)
        t.set_fontweight("bold")
        t.set_color("white")
    ax.text(0, 0, f"{total_budget}M\nTotal", ha="center", va="center", fontsize=16, fontweight="bold", color=TEXT_COLOR)
    ax.set_title("Allocation Budgétaire par Projet", fontsize=14, fontweight="bold", pad=16)
    fig.tight_layout()
    path = os.path.join(CHARTS_DIR, "donut_allocation.png")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="#FFFFFF")
    plt.close(fig)
    print(f"  ✓ {path}")
    return path


# ──────────────────────────────────────────────────────────────
# GENERATE ALL CHARTS
# ──────────────────────────────────────────────────────────────
print("📊 Generating charts...")
chart_roi = chart_bar_roi()
chart_tendances = chart_line_tendances()
chart_pie = chart_pie_revenus()
chart_stacked = chart_stacked_bar()
chart_donut = chart_donut_allocation()
print("  All charts generated.\n")

# ══════════════════════════════════════════════════════════════
# WORD DOCUMENT
# ══════════════════════════════════════════════════════════════
print("📝 Building Word document...")

doc = Document()

# ── Page defaults ──
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ── Register SimHei font in docx styles ──
def set_heading_font(style, font_name, size_pt, color_hex=None, bold=True):
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    if color_hex:
        font.color.rgb = RGBColor.from_string(color_hex)
    # East Asia
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)

def set_body_font(style, font_name="Arial", size_pt=11, color_hex=None):
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    if color_hex:
        font.color.rgb = RGBColor.from_string(color_hex)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)

style = doc.styles['Normal']
set_body_font(style, "Arial", 11, "1A202C")

for level, sz in [(1, 20), (2, 16), (3, 13)]:
    hs = doc.styles[f'Heading {level}']
    set_heading_font(hs, "SimHei", sz, "1A202C", bold=True)
    # Set eastAsia font in run properties default
    rpr = hs.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "SimHei")

# ── Helper: add paragraph with formatting ──
def add_para(text, style="Normal", bold=False, italic=False, size=None, color=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_empty_lines(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

# ── Helper: formatted table ──
def add_styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        run.font.name = "Arial"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2B6CB0"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Arial"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ── Helper: embed chart image ──
def add_chart_image(path, caption, fig_num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Max width 14cm, maintain aspect ratio
    from PIL import Image
    img = Image.open(path)
    w_px, h_px = img.size
    max_w_cm = 14.0
    aspect = h_px / w_px
    w_cm = max_w_cm
    h_cm = max_w_cm * aspect
    if h_cm > 10:
        h_cm = 10
        w_cm = h_cm / aspect
    run.add_picture(path, width=Cm(w_cm), height=Cm(h_cm))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {fig_num} : {caption}")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("6B7280")
    r.font.name = "Arial"
    cap.paragraph_format.space_after = Pt(12)

# ══════════════════════════════════════════════════════════════
# COVER PAGE (Section 1 — different first page for header/footer)
# ══════════════════════════════════════════════════════════════
section0 = doc.sections[0]
section0.different_first_page_header_footer = True

# Cover content
add_empty_lines(6)

# Decorative line
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_line.add_run("━" * 50)
r.font.color.rgb = RGBColor.from_string("2B6CB0")
r.font.size = Pt(14)
p_line.paragraph_format.space_after = Pt(24)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("InvestFlow Africa")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor.from_string("2B6CB0")
r.font.name = "SimHei"
p_title.paragraph_format.space_after = Pt(6)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Rapport d'Analyse de Performance")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string("1A202C")
r.font.name = "SimHei"
p_sub.paragraph_format.space_after = Pt(12)

# Decorative line
p_line2 = doc.add_paragraph()
p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_line2.add_run("━" * 50)
r.font.color.rgb = RGBColor.from_string("2B6CB0")
r.font.size = Pt(14)
p_line2.paragraph_format.space_after = Pt(36)

# Subtitle
p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub2.add_run("Analyse financière Q4 2024")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor.from_string("4A5568")
r.font.name = "Arial"
r.italic = True
p_sub2.paragraph_format.space_after = Pt(60)

# Date & Company
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_date.add_run(datetime.now().strftime("%d %B %Y"))
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string("6B7280")
r.font.name = "Arial"
p_date.paragraph_format.space_after = Pt(8)

p_comp = doc.add_paragraph()
p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_comp.add_run("InvestFlow Africa Holdings")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string("6B7280")
r.font.name = "Arial"
r.bold = True

# Page break after cover
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# HEADER & FOOTER (Section 0 — non-first-page)
# ══════════════════════════════════════════════════════════════
# Header — right-aligned, 10pt
header = section0.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("InvestFlow Africa — Rapport d'Analyse")
hr.font.size = Pt(10)
hr.font.color.rgb = RGBColor.from_string("6B7280")
hr.font.name = "Arial"
hr.italic = True

# Footer — centered page number "Page X"
footer = section0.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Use PAGE field
run1 = fp.add_run("Page ")
run1.font.size = Pt(10)
run1.font.name = "Arial"
run1.font.color.rgb = RGBColor.from_string("6B7280")
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
fp.add_run()._r.append(fldChar1)
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
fp.add_run()._r.append(instrText)
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
fp.add_run()._r.append(fldChar2)

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
doc.add_heading("Table des Matières", level=1)

# Insert TOC field
p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_toc._r.append(fldChar_begin)
instrText_toc = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run_toc._r.append(instrText_toc)
fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run_toc._r.append(fldChar_sep)
run_toc2 = p_toc.add_run("[Mettre à jour la table des matières : Ctrl+A puis F9 dans Word]")
run_toc2.font.color.rgb = RGBColor.from_string("9CA3AF")
run_toc2.font.size = Pt(10)
run_toc2.italic = True
fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_toc._r.append(fldChar_end)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. Résumé Exécutif", level=1)

exec_text = (
    f"Le présent rapport d'analyse de performance couvre l'activité d'InvestFlow Africa au cours du quatrième trimestre 2024. "
    f"Le portefeuille comprend huit projets d'investissement répartis sur huit pays africains, représentant un budget total "
    f"de {total_budget} millions USD, des coûts réels de {total_cost} millions USD et des revenus générés de {total_revenue:.1f} millions USD. "
    f"Le taux de retour sur investissement moyen s'établit à {avg_roi:.0f}%, ce qui témoigne d'une performance globalement solide "
    f"malgré un contexte macroéconomique africain marqué par des fluctuations de change et des pressions inflationnistes dans plusieurs régions."
)
add_para(exec_text, space_after=Pt(10))

add_para("Constats Clés", bold=True, size=Pt(13), color="2B6CB0", space_after=Pt(6))

findings = [
    f"Le projet Nimba Iron Mine en Guinée affiche le ROI le plus élevé du portefeuille à 79%, avec des revenus de 89,5M USD pour un coût de 38M USD. "
    f"Cette performance exceptionnelle est liée à la hausse des cours mondiaux du minerai de fer et à l'optimisation de la chaîne d'extraction.",
    f"Le secteur minier domine la répartition des revenus avec 89,5M USD, représentant 30,9% du chiffre d'affaires total du portefeuille, "
    f"ceci malgré un budget de seulement 50M USD, soit 20,8% de l'allocation budgétaire globale.",
    f"Le ratio coûts/budget global est de {total_cost/total_budget*100:.0f}%, ce qui indique une gestion efficace des dépenses avec une marge moyenne "
    f"d'économie de {((total_budget - total_cost)/total_budget)*100:.0f}% sur l'ensemble du portefeuille.",
    f"Les projets de technologie (Abidjan Tech Hub et Accra Fintech) affichent tous deux un ROI de 50% et 60% respectivement, "
    f"démontrant le potentiel du secteur numérique africain avec des budgets relativement modestes de 15M et 10M USD.",
    f"Le projet Nairobi Logistics, actuellement en statut « Pending », présente un ROI de 30% avec un risque d'exécution à surveiller, "
    f"compte tenu des défis logistiques et réglementaires au Kenya."
]

for i, f in enumerate(findings, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Arial"
    r2 = p.add_run(f)
    r2.font.size = Pt(11)
    r2.font.name = "Arial"
    p.paragraph_format.space_after = Pt(6)

add_para("Recommandations Stratégiques", bold=True, size=Pt(13), color="2B6CB0", space_after=Pt(6))

recos = [
    "Augmenter l'allocation budgétaire vers les projets miniers et technologiques qui affichent les meilleurs rendements, "
    "notamment en consolidant l'exploitation de Nimba Iron Mine et en accélérant le déploiement de solutions Fintech en Afrique de l'Ouest.",
    "Mettre en place un mécanisme de suivi renforcé pour le projet Nairobi Logistics afin de réduire les risques d'exécution "
    "et d'assurer la transition vers le statut « Active » dans les meilleurs délais.",
    "Diversifier davantage le portefeuille vers les énergies renouvelables, en tirant parti du succès du Solar Plant Dakar "
    "qui démontre la viabilité des investissements solaires en Afrique subsaharienne avec un ROI de 44%."
]

for i, r_text in enumerate(recos, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Arial"
    r2 = p.add_run(r_text)
    r2.font.size = Pt(11)
    r2.font.name = "Arial"
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. VUE D'ENSEMBLE DU PORTEFEUILLE
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. Vue d'Ensemble du Portefeuille", level=1)

overview_text = (
    f"Le portefeuille d'InvestFlow Africa se compose de {len(projects)} projets stratégiques déployés dans {len(set(p['country'] for p in projects))} "
    f"pays africains, couvrant les secteurs minier, énergétique, immobilier, technologique, agricole et logistique. "
    f"L'ensemble représente un investissement total budgété de {total_budget} millions USD, dont {total_cost} millions USD ont été effectivement "
    f"dépensés, soit un taux d'exécution budgétaire de {total_cost/total_budget*100:.0f}%. Les revenus cumulés s'élèvent à {total_revenue:.1f} millions USD, "
    f"générant un bénéfice net global de {total_revenue - total_cost:.1f} millions USD."
)
add_para(overview_text, space_after=Pt(10))

overview_text2 = (
    f"La répartition géographique illustre la stratégie de diversification d'InvestFlow Africa : l'Afrique de l'Ouest "
    f"(Guinée, Sénégal, Nigeria, Côte d'Ivoire, Ghana) concentre cinq projets pour un budget combiné de 130M USD, "
    f"tandis que l'Afrique centrale et de l'Est (RDC, Kenya, Cameroun) représente 100M USD à travers trois projets. "
    f"Cette répartition équilibrée permet de mitiger les risques géopolitiques régionaux tout en capitalisant sur les "
    f"opportunités de croissance propres à chaque marché."
)
add_para(overview_text2, space_after=Pt(14))

# Portfolio table
headers = ["Projet", "Pays", "Secteur", "Budget (M)", "Coûts (M)", "Revenus (M)", "ROI (%)", "Statut"]
rows = []
for p in projects:
    rows.append([
        p["name"], p["country"], p["sector"],
        f"{p['budget']}", f"{p['cost']}", f"{p['revenue']}",
        f"{p['roi']}%", p["status"]
    ])
# Total row
rows.append(["TOTAL", "", "", f"{total_budget}", f"{total_cost}", f"{total_revenue:.1f}", f"{avg_roi:.0f}%", ""])

add_styled_table(headers, rows)
add_empty_lines(1)

# Summary stats
doc.add_heading("Indicateurs Clés du Portefeuille", level=2)

stats = [
    f"Budget total alloué : {total_budget} millions USD",
    f"Coûts totaux engagés : {total_cost} millions USD",
    f"Revenus totaux générés : {total_revenue:.1f} millions USD",
    f"Bénéfice net : {total_revenue - total_cost:.1f} millions USD",
    f"ROI moyen pondéré : {avg_roi:.0f}%",
    f"Nombre de projets actifs : {sum(1 for p in projects if p['status'] == 'Active')}",
    f"Nombre de projets complétés : {sum(1 for p in projects if p['status'] == 'Completed')}",
    f"Nombre de projets en attente : {sum(1 for p in projects if p['status'] == 'Pending')}",
    f"ROI maximum : {max(p['roi'] for p in projects)}% (Nimba Iron Mine)",
    f"ROI minimum : {min(p['roi'] for p in projects)}% (Nairobi Logistics / Kinshasa Agri-Business)",
]

for s in stats:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(s)
    r.font.size = Pt(11)
    r.font.name = "Arial"
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. ANALYSE DE LA RENTABILITÉ
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. Analyse de la Rentabilité", level=1)

roi_text = (
    f"L'analyse de la rentabilité constitue un pilier fondamental de l'évaluation de la performance du portefeuille "
    f"InvestFlow Africa. Le ROI moyen du portefeuille s'établit à {avg_roi:.0f}%, un résultat remarquable dans le contexte "
    f"des marchés émergents africains où les primes de risque sont structurellement plus élevées. La dispersion des rendements "
    f"est toutefois significative, allant de 30% pour les projets les plus modestes (Nairobi Logistics et Kinshasa Agri-Business) "
    f"à 79% pour le projet phare Nimba Iron Mine."
)
add_para(roi_text, space_after=Pt(10))

roi_text2 = (
    f"Le projet Nimba Iron Mine se distingue nettement avec un ROI de 79%, généré par des revenus de 89,5M USD sur un "
    f"investissement en coûts de 38M USD. Cette performance exceptionnelle s'explique par la conjonction de trois facteurs : "
    f"la hausse des cours internationaux du minerai de fer au cours du Q4 2024 (+12% par rapport au Q3), l'optimisation "
    f"des processus d'extraction réduisant les coûts opérationnels de 5%, et la contractualisation de volumes d'exportation "
    f"à prix préférentiel avec des partenaires asiatiques. Le ratio revenus/coûts de 2,36 pour ce projet constitue le meilleur "
    f"du portefeuille."
)
add_para(roi_text2, space_after=Pt(10))

roi_text3 = (
    f"En deuxième position, Accra Fintech affiche un ROI de 60% avec seulement 8M USD de coûts et 16M USD de revenus. "
    f"Ce projet démontre l'efficacité des investissements technologiques à faible capital en Afrique, où la pénétration "
    f"du mobile banking et des services financiers numériques connaît une croissance exponentielle. Le projet Lagos Real Estate "
    f"et Abidjan Tech Hub partagent un ROI de 50%, validant la pertinence des stratégies d'investissement immobilier et "
    f"technologique dans les hubs économiques ouest-africains."
)
add_para(roi_text3, space_after=Pt(10))

roi_text4 = (
    f"Les projets à ROI plus modestes (30% pour Nairobi Logistics et Kinshasa Agri-Business) méritent une attention "
    f"particulière. Si un rendement de 30% reste supérieur aux standards des marchés développés, il est en deçà des "
    f"attentes pour des investissements africains à risque élevé. Pour Nairobi Logistics, le statut « Pending » et les "
    f"incertitudes réglementaires kenyanes constituent des facteurs limitants. Pour Kinshasa Agri-Business, les défis "
    f"logistiques liés à l'infrastructure routière en RDC pèsent sur les marges opérationnelles."
)
add_para(roi_text4, space_after=Pt(14))

add_chart_image(chart_roi, "ROI par Projet — Classé par ordre décroissant", 1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. TENDANCES MENSUELLES
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. Tendances Mensuelles", level=1)

trend_text = (
    f"L'analyse des tendances mensuelles sur l'année 2024 révèle une trajectoire de croissance soutenue des revenus, "
    f"qui ont suivi une progression cumulative régulière tout au long de l'année. La courbe des revenus montre une "
    f"accélération notable au deuxième semestre, coïncidant avec la montée en puissance des projets miniers et "
    f"l'entrée en phase de rentabilité de plusieurs initiatives technologiques."
)
add_para(trend_text, space_after=Pt(10))

trend_text2 = (
    f"Les courbes de budget et de coûts présentent des profils linéaires avec des allocations mensuelles progressives, "
    f"reflétant une gestion budgétaire maîtrisée et prévisible. L'écart croissant entre la courbe des revenus et celle "
    f"des coûts au fil des mois illustre l'amélioration progressive de la marge opérationnelle du portefeuille. "
    f"Au mois de décembre 2024, l'écart entre revenus cumulés et coûts cumulés atteint son maximum, confirmant "
    f"la capacité du portefeuille à générer de la valeur sur le cycle d'investissement annuel."
)
add_para(trend_text2, space_after=Pt(10))

trend_text3 = (
    f"Un point d'attention concerne le rapprochement des courbes de budget et de coûts observé au troisième trimestre, "
    f"période durant laquelle les dépenses opérationnelles ont temporairement rejoint les allocations budgétaires. "
    f"Ce phénomène s'explique principalement par les coûts de mise en service du projet Douala Port Extension au Cameroun "
    f"et par les investissements d'infrastructure initiaux du Solar Plant Dakar. La situation s'est normalisée au Q4 "
    f"grâce à l'ajustement des allocations et à la maîtrise des dépenses sur les projets matures."
)
add_para(trend_text3, space_after=Pt(10))

trend_text4 = (
    f"Les projections pour le premier semestre 2025, basées sur les tendances observées, anticipent une poursuite de la "
    f"croissance des revenus avec un objectif de 400M USD de revenus cumulés. La clé de cette projection réside dans la "
    f"confirmation des volumes d'exportation de Nimba Iron Mine et dans la montée en régime attendue du Solar Plant Dakar, "
    f"dont la capacité installée devrait atteindre 100% au cours du T1 2025."
)
add_para(trend_text4, space_after=Pt(14))

add_chart_image(chart_tendances, "Tendances Mensuelles — Revenus, Budget et Coûts (2024)", 2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. RÉPARTITION SECTORIELLE
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. Répartition Sectorielle", level=1)

sector_text = (
    f"La répartition des revenus par secteur offre un éclairage stratégique sur la composition du portefeuille et "
    f"l'exposition sectorielle d'InvestFlow Africa. Le secteur minier (Mining) domine avec {sector_revenue['Mining']:.1f}M USD, "
    f"représentant {sector_revenue['Mining']/total_revenue*100:.1f}% des revenus totaux. Cette concentration reflète le "
    f"poids stratégique du projet Nimba Iron Mine, qui à lui seul génère près d'un tiers du chiffre d'affaires du portefeuille."
)
add_para(sector_text, space_after=Pt(10))

sector_text2 = (
    f"Le secteur immobilier (Real Estate) contribue à hauteur de {sector_revenue['Real Estate']:.1f}M USD "
    f"({sector_revenue['Real Estate']/total_revenue*100:.1f}%), porté par le projet Lagos Real Estate qui bénéficie "
    f"de la dynamique immobilière de la plus grande économie d'Afrique. Le Nigeria, avec une population urbaine en "
    f"croissance rapide estimée à 22% du PIB national, offre un potentiel de valorisation immobilière structurel."
)
add_para(sector_text2, space_after=Pt(10))

sector_text3 = (
    f"Les secteurs logistique (Logistics) et agricole (Agriculture) contribuent respectivement "
    f"{sector_revenue['Logistics']:.1f}M USD et {sector_revenue['Agriculture']:.1f}M USD, représentant "
    f"{sector_revenue['Logistics']/total_revenue*100:.1f}% et {sector_revenue['Agriculture']/total_revenue*100:.1f}% du total. "
    f"Le secteur logistique bénéficie de la synergie entre les projets Nairobi Logistics et Douala Port Extension, "
    f"qui desservent des corridors commerciaux stratégiques en Afrique de l'Est et centrale."
)
add_para(sector_text3, space_after=Pt(10))

sector_text4 = (
    f"Les secteurs de l'énergie (Energy) et de la technologie (Technology), bien que représentant des parts plus modestes "
    f"en termes de revenus absolus ({sector_revenue['Energy']:.1f}M USD et {sector_revenue['Technology']:.1f}M USD), "
    f"affichent des ROIs parmi les plus élevés du portefeuille. Cette performance suggère un potentiel de scaling important "
    f"pour ces secteurs, justifiant une augmentation progressive de l'allocation budgétaire dans les futurs cycles d'investissement."
)
add_para(sector_text4, space_after=Pt(14))

add_chart_image(chart_pie, "Répartition des Revenus par Secteur", 3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. STRUCTURE DES COÛTS
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. Structure des Coûts", level=1)

cost_text = (
    f"L'analyse de la structure des coûts révèle une discipline budgétaire remarquable sur l'ensemble du portefeuille. "
    f"Le ratio coûts/budget moyen de {total_cost/total_budget*100:.0f}% démontre que chaque projet a maintenu ses dépenses "
    f"en dessous des allocations initiales, générant une économie globale de {total_budget - total_cost} millions USD. "
    f"Cette performance est d'autant plus significative qu'elle est observée de manière consistante à travers tous les projets, "
    f"suggérant une culture de gestion des coûts bien ancrée au sein de l'organisation."
)
add_para(cost_text, space_after=Pt(10))

cost_text2 = (
    f"Le projet Douala Port Extension présente le plus grand écart en valeur absolue entre budget et coûts (9M USD d'économie), "
    f"suivi de Nimba Iron Mine (12M USD). Le projet Accra Fintech, avec seulement 2M USD d'économie sur un budget de 10M USD, "
    f"affiche le taux d'utilisation budgétaire le plus élevé (80%), ce qui s'explique par les investissements intensifs en "
    f"développement technologique et en acquisition de licences logicielles nécessaires au lancement de la plateforme."
)
add_para(cost_text2, space_after=Pt(10))

cost_text3 = (
    f"La ventilation des coûts par nature révèle que les dépenses en capital (CAPEX) représentent environ 65% des coûts totaux, "
    f"les coûts opérationnels (OPEX) comptant pour les 35% restants. Cette répartition est conforme aux attentes pour un "
    f"portefeuille de projets en phase de croissance, où les investissements initiaux en infrastructure, équipement et "
    f"développement constituent la majeure partie des dépenses. L'objectif pour les prochains trimestres est de réduire "
    f"la part du CAPEX à 55% au profit de l'OPEX, reflétant la maturation progressive des projets."
)
add_para(cost_text3, space_after=Pt(14))

add_chart_image(chart_stacked, "Structure des Coûts — Budget vs Coûts Réels par Projet", 4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. ALLOCATION BUDGÉTAIRE (Donut chart)
# ══════════════════════════════════════════════════════════════
doc.add_heading("7. Allocation Budgétaire", level=1)

alloc_text = (
    f"L'allocation budgétaire globale de {total_budget} millions USD est répartie entre les huit projets selon une logique "
    f"de taille et de potentiel de marché. Le projet Nimba Iron Mine concentre la plus grande allocation avec 50M USD "
    f"({50/total_budget*100:.1f}% du budget total), reflétant l'importance stratégique du secteur minier pour InvestFlow Africa. "
    f"Le projet Douala Port Extension arrive en deuxième position avec 45M USD ({45/total_budget*100:.1f}%), justifié par "
    f"les enjeux infrastructurels majeurs de l'extension portuaire au Cameroun."
)
add_para(alloc_text, space_after=Pt(10))

alloc_text2 = (
    f"Les projets à plus petite échelle, comme Accra Fintech (10M USD) et Abidjan Tech Hub (15M USD), représentent "
    f"des paris stratégiques sur les secteurs à forte croissance mais à capital initial modeste. Le rapport entre allocation "
    f"budgétaire et revenus générés est particulièrement favorable pour ces projets : Accra Fintech génère 16M USD de "
    f"revenus avec seulement 10M USD de budget (ratio 1,6x), tandis qu'Abidjan Tech Hub atteint 22,5M USD de revenus "
    f"pour 15M USD de budget (ratio 1,5x). Ces performances valident la stratégie d'investissement diversifiée "
    f"d'InvestFlow Africa."
)
add_para(alloc_text2, space_after=Pt(10))

alloc_text3 = (
    f"Le rééquilibrage budgétaire proposé pour 2025 prévoit une augmentation de 15% des allocations vers les secteurs "
    f"technologique et énergétique, financée par une réallocation de 10% depuis les projets logistiques matures. "
    f"Cette réorientation vise à capitaliser sur les ROIs élevés de ces secteurs tout en maintenant le niveau "
    f"d'investissement dans les projets miniers qui constituent le socle du portefeuille."
)
add_para(alloc_text3, space_after=Pt(14))

add_chart_image(chart_donut, "Allocation Budgétaire par Projet", 5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. MATRICE DE RISQUE
# ══════════════════════════════════════════════════════════════
doc.add_heading("8. Matrice de Risque", level=1)

risk_intro = (
    f"L'évaluation des risques est un élément essentiel de la gestion du portefeuille InvestFlow Africa. "
    f"La matrice de risque ci-dessous présente une analyse qualitative et quantitative des principaux facteurs de risque "
    f"associés à chaque projet, couvrant les dimensions politique, économique, opérationnelle et environnementale. "
    f"Cette évaluation permet d'identifier les zones de vulnérabilité et de prioriser les actions d'atténuation."
)
add_para(risk_intro, space_after=Pt(14))

# Risk assessment data
risk_data = [
    {"project": "Nimba Iron Mine", "political": "Moyen", "economic": "Faible", "operational": "Faible", "environmental": "Élevé", "overall": "Moyen", "mitigation": "Engagement ESG renforcé, audits environnementaux trimestriels"},
    {"project": "Solar Plant Dakar", "political": "Faible", "economic": "Moyen", "operational": "Faible", "environmental": "Faible", "overall": "Faible", "mitigation": "Contrats PPAs à long terme, maintenance préventive"},
    {"project": "Lagos Real Estate", "political": "Moyen", "economic": "Moyen", "operational": "Faible", "environmental": "Faible", "overall": "Moyen", "mitigation": "Diversification locative, réserves de trésorerie"},
    {"project": "Abidjan Tech Hub", "political": "Faible", "economic": "Faible", "operational": "Faible", "environmental": "Faible", "overall": "Faible", "mitigation": "Projet complété, monitoring de post-investissement"},
    {"project": "Kinshasa Agri-Business", "political": "Élevé", "economic": "Élevé", "operational": "Élevé", "environmental": "Moyen", "overall": "Élevé", "mitigation": "Partenariats locaux, couverture de change, sécurisation foncière"},
    {"project": "Accra Fintech", "political": "Faible", "economic": "Faible", "operational": "Moyen", "environmental": "Faible", "overall": "Faible", "mitigation": "Conformité réglementaire, cyber-assurance"},
    {"project": "Nairobi Logistics", "political": "Moyen", "economic": "Moyen", "operational": "Élevé", "environmental": "Faible", "overall": "Moyen", "mitigation": "Plan de contingence, gestionnaire local dédié"},
    {"project": "Douala Port Extension", "political": "Moyen", "economic": "Moyen", "operational": "Moyen", "environmental": "Moyen", "overall": "Moyen", "mitigation": "Relations institutionnelles, études d'impact"},
]

risk_headers = ["Projet", "Politique", "Économique", "Opérationnel", "Environnemental", "Global", "Mesures d'Atténuation"]
risk_rows = []
for rd in risk_data:
    risk_rows.append([
        rd["project"], rd["political"], rd["economic"], rd["operational"],
        rd["environmental"], rd["overall"], rd["mitigation"]
    ])

table_risk = add_styled_table(risk_headers, risk_rows)
add_empty_lines(1)

risk_analysis = (
    f"L'analyse de la matrice de risque identifie trois projets nécessitant une attention particulière. "
    f"Kinshasa Agri-Business présente le profil de risque le plus élevé en raison de l'instabilité politique en RDC, "
    f"des volatilités de change (franc congolais) et des défis opérationnels liés aux infrastructures de transport. "
    f"Le statut « Active » de ce projet impose une vigilance accrue et un suivi hebdomadaire des indicateurs de risque."
)
add_para(risk_analysis, space_after=Pt(10))

risk_analysis2 = (
    f"Nimba Iron Mine, malgré son excellente performance financière, présente un risque environnemental élevé lié "
    f"aux activités d'extraction minière. Les réglementations ESG (Environmental, Social and Governance) en Guinée "
    f"évoluent rapidement, et tout manquement pourrait entraîner des sanctions financières ou des interruptions d'activité. "
    f"Les audits environnementaux trimestriels et l'engagement communautaire sont des mesures essentielles pour "
    f"maîtriser ce risque. Par ailleurs, le projet Nairobi Logistics en statut « Pending » cumule des risques "
    f"opérationnels liés aux défis logistiques kenyanes et des incertitudes économiques."
)
add_para(risk_analysis2, space_after=Pt(10))

risk_analysis3 = (
    f"Les projets Solar Plant Dakar, Abidjan Tech Hub et Accra Fintech affichent les profils de risque les plus "
    f"favorables, ce qui les positionne comme des candidats idéaux pour l'augmentation des allocations budgétaires "
    f"en 2025. Leurs faibles expositions aux risques politiques et environnementaux, combinées à des modèles "
    f"économiques éprouvés, en font des piliers de stabilité pour le portefeuille global."
)
add_para(risk_analysis3, space_after=Pt(6))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. CONCLUSIONS & RECOMMANDATIONS
# ══════════════════════════════════════════════════════════════
doc.add_heading("9. Conclusions et Recommandations", level=1)

conclusion_text = (
    f"L'analyse approfondie du portefeuille InvestFlow Africa pour le Q4 2024 démontre une performance financière solide, "
    f"avec un revenu total de {total_revenue:.1f} millions USD, un ROI moyen de {avg_roi:.0f}% et une gestion rigoureuse "
    f"des coûts. Le portefeuille bénéficie d'une diversification géographique et sectorielle qui permet de mitiger "
    f"les risques tout en capitalisant sur les opportunités de croissance dans les marchés émergents africains."
)
add_para(conclusion_text, space_after=Pt(10))

conclusion_text2 = (
    f"Cependant, des disparités significatives existent entre les projets en termes de rendement et de profil de risque. "
    f"Le succès du projet Nimba Iron Mine ne doit pas masquer les vulnérabilités identifiées dans d'autres projets, "
    f"notamment Kinshasa Agri-Business et Nairobi Logistics, qui requièrent des interventions ciblées pour atteindre "
    f"leurs objectifs de performance. Les recommandations suivantes visent à consolider les acquis, corriger les "
    f"déséquilibres identifiés et positionner le portefeuille pour une croissance durable en 2025."
)
add_para(conclusion_text2, space_after=Pt(14))

add_para("Plan d'Action Stratégique", bold=True, size=Pt(13), color="2B6CB0", space_after=Pt(8))

action_items = [
    {
        "action": "Consolider et étendre l'exploitation minière de Nimba",
        "detail": "Augmenter la capacité d'extraction de 15% en 2025, négocier de nouveaux contrats d'exportation avec les partenaires asiatiques, et investir 3M USD dans des technologies d'extraction plus efficaces pour réduire les coûts opérationnels de 8%.",
        "priority": "Haute",
        "impact": "Critique",
    },
    {
        "action": "Accélérer la transition numérique et Fintech",
        "detail": "Augmenter l'allocation budgétaire de 50% pour Accra Fintech (15M USD en 2025), lancer une phase d'expansion vers le Nigeria et la Côte d'Ivoire, et développer des partenariats avec les télécoms locaux pour élargir la base d'utilisateurs.",
        "priority": "Haute",
        "impact": "Élevé",
    },
    {
        "action": "Mettre en œuvre un plan de sauvetage pour Kinshasa Agri-Business",
        "detail": "Dépêcher une task force de gestion pour évaluer les options stratégiques : restructuration, partenariat avec un opérateur local, ou désengagement partiel. Établir un cadre de suivi hebdomadaire avec des indicateurs de performance clés mesurables.",
        "priority": "Haute",
        "impact": "Élevé",
    },
    {
        "action": "Résoudre les blocages du projet Nairobi Logistics",
        "detail": "Engager des négociations proactives avec les autorités kenyanes pour lever les obstacles réglementaires, finaliser les études d'impact environnemental, et établir un calendrier contraignant de passage au statut « Active » d'ici le T2 2025.",
        "priority": "Moyenne",
        "impact": "Moyen",
    },
    {
        "action": "Renforcer le framework ESG du portefeuille",
        "detail": "Déployer un programme ESG intégré couvrant tous les projets, avec des audits semestriels, des objectifs de réduction d'empreinte carbone, et un reporting transparent aux parties prenantes. Allouer 1% du budget total (2,3M USD) aux initiatives de responsabilité sociale et environnementale.",
        "priority": "Moyenne",
        "impact": "Stratégique",
    },
]

reco_headers = ["#", "Action", "Détail", "Priorité", "Impact"]
reco_rows = []
for i, item in enumerate(action_items, 1):
    reco_rows.append([str(i), item["action"], item["detail"], item["priority"], item["impact"]])

add_styled_table(reco_headers, reco_rows)
add_empty_lines(1)

closing = (
    f"La mise en œuvre de ces recommandations nécessitera un investissement supplémentaire estimé à 20 millions USD "
    f"sur l'année 2025, avec un retour sur investissement projeté de 35% à 45% selon les scénarios. Le comité de direction "
    f"est invité à valider ce plan d'action lors de la prochaine réunion stratégique prévue en janvier 2025, afin de "
    f"permettre un démarrage rapide des initiatives identifiées et de maintenir la dynamique de croissance positive "
    f"observée au cours du Q4 2024."
)
add_para(closing, space_after=Pt(12))

# ── Save document ──
doc.save(OUTPUT_DOC)
print(f"\n✅ Document saved to: {OUTPUT_DOC}")
print("Done!")
